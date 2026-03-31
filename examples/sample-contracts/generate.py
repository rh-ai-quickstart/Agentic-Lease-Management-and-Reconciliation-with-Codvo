"""
Generate 30 realistic aircraft leasing intake documents (DOCX) for LeasingOps testing.
Full formatting: cover pages, logos, headers/footers, color scheme, table styling.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import datetime, os, io, copy

OUT = os.path.dirname(os.path.abspath(__file__))

# ─── Brand palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x2B, 0x55)   # deep navy — section headings
BLUE   = RGBColor(0x1A, 0x56, 0x9E)   # corporate blue — sub-headings
GOLD   = RGBColor(0xC8, 0x9A, 0x2C)   # gold accent — cover badge border
LIGHT  = RGBColor(0xE8, 0xEF, 0xF8)   # light blue — table header fill
MID    = RGBColor(0xF4, 0xF7, 0xFC)   # very light — table alt-row fill
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x66, 0x66, 0x66)
RED    = RGBColor(0xCC, 0x00, 0x00)

# Hex helpers for oxml (need RRGGBB string without #)
def hex_color(r, g, b): return f"{r:02X}{g:02X}{b:02X}"

NAVY_HEX  = "0D2B55"
BLUE_HEX  = "1A569E"
LIGHT_HEX = "E8EFF8"
MID_HEX   = "F4F7FC"
GOLD_HEX  = "C89A2C"
WHITE_HEX = "FFFFFF"
GRAY_HEX  = "666666"

# ─── Logo generator ───────────────────────────────────────────────────────────

def make_logo(name: str, primary: tuple, accent: tuple, width=320, height=80) -> io.BytesIO:
    """Generate a professional company logo PNG (initials badge + full name)."""
    img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Badge circle
    r = 34
    cx, cy = 40, height // 2
    draw.ellipse([cx-r, cy-r, cx+r, cy+r], fill=primary)

    # Initials in badge
    initials = "".join(w[0].upper() for w in name.split() if w[0].isalpha())[:3]
    try:
        font_bold = ImageFont.truetype("/System/Library/Fonts/HelveticaNeue.ttc", 22)
        font_name = ImageFont.truetype("/System/Library/Fonts/HelveticaNeue.ttc", 17)
        font_sub  = ImageFont.truetype("/System/Library/Fonts/HelveticaNeue.ttc", 11)
    except Exception:
        font_bold = ImageFont.load_default()
        font_name = font_bold
        font_sub  = font_bold

    # Center initials in badge
    bbox = draw.textbbox((0, 0), initials, font=font_bold)
    tw, th = bbox[2]-bbox[0], bbox[3]-bbox[1]
    draw.text((cx - tw//2, cy - th//2 - 2), initials, fill=(255,255,255,255), font=font_bold)

    # Company name next to badge
    tx = cx + r + 12
    # Split name into two lines if long
    words = name.split()
    line1 = " ".join(words[:len(words)//2 + (len(words)%2)]) if len(words) > 2 else name
    line2 = " ".join(words[len(words)//2 + (len(words)%2):]) if len(words) > 2 else ""
    draw.text((tx, cy - 16), line1, fill=primary+(255,), font=font_name)
    if line2:
        draw.text((tx, cy + 4), line2, fill=accent+(255,), font=font_sub)

    # Gold bottom rule
    draw.rectangle([cx-r, cy+r+3, width-4, cy+r+6], fill=accent)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


LESSOR_LOGOS = {}
LESSEE_LOGOS = {}

def ensure_logos(lessor: str, lessee: str):
    if lessor not in LESSOR_LOGOS:
        LESSOR_LOGOS[lessor] = make_logo(lessor, (13,43,85), (200,154,44))
    if lessee not in LESSEE_LOGOS:
        LESSEE_LOGOS[lessee] = make_logo(lessee, (26,86,158), (13,43,85))


# ─── oxml helpers for rich formatting ─────────────────────────────────────────

def set_cell_bg(cell, hex_col: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col)
    tcPr.append(shd)

def set_cell_bold_white(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = WHITE
        if not para.runs and para.text:
            run = para.runs[0] if para.runs else para.add_run(para.text)
            run.bold = True
            run.font.color.rgb = WHITE

def add_horiz_rule(doc):
    """Add a thin navy horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def set_para_font(para, size=10, bold=False, color=None, italic=False, name="Calibri"):
    for run in para.runs:
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

def set_run_font(run, size=10, bold=False, color=None, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_table_styled(doc, headers, rows, col_widths=None, alt_rows=True):
    """Full-styled table: navy header row, alt-row shading, gridlines."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, NAVY_HEX)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg = MID_HEX if (alt_rows and r_idx % 2 == 1) else WHITE_HEX
        for c_idx, val in enumerate(row_data):
            cell = cells[c_idx]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(1)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths[:len(row.cells)]):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        r = p.runs[0]
        r.font.name = "Calibri"
        if level == 1:
            r.font.color.rgb = NAVY
            r.font.size = Pt(16)
            r.bold = True
        elif level == 2:
            r.font.color.rgb = NAVY
            r.font.size = Pt(12)
            r.bold = True
        elif level == 3:
            r.font.color.rgb = BLUE
            r.font.size = Pt(10)
            r.bold = True
    return p

def body(doc, text, indent=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.italic = italic
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def page_break(doc):
    doc.add_page_break()

def sig_block(doc, parties):
    add_horiz_rule(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("IN WITNESS WHEREOF the parties have executed this document as of the date first written above.")
    r.font.name = "Calibri"; r.font.size = Pt(10)
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=len(parties))
    table.style = "Table Grid"
    for i, (label, name, title, date) in enumerate(parties):
        for row_i, val in enumerate([label, "By: _________________________", f"Name: {name}", f"Title: {title}", f"Date: {date}"]):
            cell = table.rows[row_i].cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if row_i == 0:
                run.bold = True
                run.font.color.rgb = NAVY
                set_cell_bg(cell, LIGHT_HEX)
    doc.add_paragraph()


# ─── Header / Footer ─────────────────────────────────────────────────────────

def add_header_footer(doc, doc_ref: str, doc_type: str, lessor: str):
    """Add professional header and footer to all sections."""
    section = doc.sections[0]
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 3, width=Inches(6.5))
    htable.style = "Table Grid"
    # Remove table borders
    tbl = htable._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Left: lessor name
    lc = htable.rows[0].cells[0]
    lc.paragraphs[0].clear()
    r = lc.paragraphs[0].add_run(lessor)
    r.font.name = "Calibri"; r.font.size = Pt(8); r.bold = True
    r.font.color.rgb = NAVY

    # Center: doc type
    mc = htable.rows[0].cells[1]
    mc.paragraphs[0].clear()
    mc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = mc.paragraphs[0].add_run(doc_type.upper())
    r2.font.name = "Calibri"; r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY

    # Right: ref
    rc = htable.rows[0].cells[2]
    rc.paragraphs[0].clear()
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = rc.paragraphs[0].add_run(doc_ref)
    r3.font.name = "Calibri"; r3.font.size = Pt(8)
    r3.font.color.rgb = GRAY

    # Header bottom rule
    add_horiz_rule(header)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    ftable = footer.add_table(1, 3, width=Inches(6.5))
    ftable.style = "Table Grid"
    # Remove borders same way
    ftbl = ftable._tbl
    ftblPr = ftbl.find(qn("w:tblPr"))
    if ftblPr is None:
        ftblPr = OxmlElement("w:tblPr")
        ftbl.insert(0, ftblPr)
    ftblBorders = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        ftblBorders.append(b)
    ftblPr.append(ftblBorders)

    # Add top rule to footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1"); top.set(qn("w:color"), NAVY_HEX)
    pBdr.append(top); pPr.append(pBdr)

    # Left: confidential
    fc = ftable.rows[0].cells[0]
    fc.paragraphs[0].clear()
    r4 = fc.paragraphs[0].add_run("CONFIDENTIAL — PRIVILEGED & PROPRIETARY")
    r4.font.name = "Calibri"; r4.font.size = Pt(7)
    r4.font.color.rgb = GRAY; r4.italic = True

    # Center: page number using Word field
    pc = ftable.rows[0].cells[1]
    pc.paragraphs[0].clear()
    pc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = pc.paragraphs[0].add_run("Page ")
    r5.font.name = "Calibri"; r5.font.size = Pt(8); r5.font.color.rgb = GRAY
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run_pg = OxmlElement("w:r")
    run_pg.append(fldChar1)
    run_pg2 = OxmlElement("w:r")
    run_pg2.append(instrText)
    run_pg3 = OxmlElement("w:r")
    run_pg3.append(fldChar2)
    pc.paragraphs[0]._p.append(run_pg)
    pc.paragraphs[0]._p.append(run_pg2)
    pc.paragraphs[0]._p.append(run_pg3)
    r6 = pc.paragraphs[0].add_run(" of ")
    r6.font.name = "Calibri"; r6.font.size = Pt(8); r6.font.color.rgb = GRAY
    # NUMPAGES field
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    iT = OxmlElement("w:instrText"); iT.text = " NUMPAGES "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    rn1 = OxmlElement("w:r"); rn1.append(f1)
    rn2 = OxmlElement("w:r"); rn2.append(iT)
    rn3 = OxmlElement("w:r"); rn3.append(f2)
    pc.paragraphs[0]._p.append(rn1)
    pc.paragraphs[0]._p.append(rn2)
    pc.paragraphs[0]._p.append(rn3)

    # Right: date
    dc = ftable.rows[0].cells[2]
    dc.paragraphs[0].clear()
    dc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r7 = dc.paragraphs[0].add_run(datetime.date.today().strftime("%d %b %Y"))
    r7.font.name = "Calibri"; r7.font.size = Pt(8); r7.font.color.rgb = GRAY


# ─── Cover page ──────────────────────────────────────────────────────────────

def add_cover_page(doc, doc_type: str, doc_ref: str, ac: dict,
                   subtitle: str, parties_table: list, status="EXECUTION COPY"):
    """Full-bleed cover page: logo bar, title, aircraft summary, parties table, status badge."""
    ensure_logos(ac["lessor"], ac["lessee"])

    # ── Navy top banner ──
    banner = doc.add_paragraph()
    banner.paragraph_format.space_before = Pt(0)
    banner.paragraph_format.space_after  = Pt(0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), NAVY_HEX)
    banner._p.get_or_add_pPr().append(shd)
    r = banner.add_run("  " + ac["lessor"].upper() + "   |   AIRCRAFT LEASING")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True
    r.font.color.rgb = WHITE

    doc.add_paragraph()

    # ── Lessor logo ──
    logo_buf = LESSOR_LOGOS[ac["lessor"]]
    logo_buf.seek(0)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_para.add_run()
    logo_run.add_picture(logo_buf, width=Inches(2.8))

    # ── Gold divider ──
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), GOLD_HEX)
    pBdr.append(bot); pPr.append(pBdr)
    hr.paragraph_format.space_after = Pt(18)

    # ── Document type ──
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_type = p_type.add_run(doc_type.upper())
    r_type.font.name = "Calibri"; r_type.font.size = Pt(11)
    r_type.font.color.rgb = BLUE; r_type.bold = True
    p_type.paragraph_format.space_after = Pt(4)

    # ── Main title ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run(subtitle)
    r_title.font.name = "Calibri"; r_title.font.size = Pt(22)
    r_title.font.color.rgb = NAVY; r_title.bold = True
    p_title.paragraph_format.space_after = Pt(6)

    # ── Reference ──
    p_ref = doc.add_paragraph()
    r_ref = p_ref.add_run(f"Reference: {doc_ref}")
    r_ref.font.name = "Calibri"; r_ref.font.size = Pt(10)
    r_ref.font.color.rgb = GRAY
    p_ref.paragraph_format.space_after = Pt(20)

    # ── Aircraft summary box ──
    ac_table = doc.add_table(rows=2, cols=5)
    ac_table.style = "Table Grid"
    labels = ["Aircraft Type", "MSN", "Registration", "Engines", "Year"]
    vals = [ac["type"], ac["msn"], ac["reg"], ac["engines"], str(ac["year"])]
    for i, (lbl, val) in enumerate(zip(labels, vals)):
        lc = ac_table.rows[0].cells[i]
        set_cell_bg(lc, NAVY_HEX)
        lc.paragraphs[0].clear()
        r = lc.paragraphs[0].add_run(lbl)
        r.font.name = "Calibri"; r.font.size = Pt(8); r.bold = True
        r.font.color.rgb = WHITE
        vc = ac_table.rows[1].cells[i]
        set_cell_bg(vc, LIGHT_HEX)
        vc.paragraphs[0].clear()
        r2 = vc.paragraphs[0].add_run(val)
        r2.font.name = "Calibri"; r2.font.size = Pt(9); r2.bold = True
        r2.font.color.rgb = NAVY
    doc.add_paragraph()

    # ── Parties table ──
    if parties_table:
        pt = doc.add_table(rows=1 + len(parties_table), cols=3)
        pt.style = "Table Grid"
        for i, h in enumerate(["Role", "Party", "Jurisdiction"]):
            c = pt.rows[0].cells[i]
            set_cell_bg(c, NAVY_HEX)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(h)
            r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True
            r.font.color.rgb = WHITE
        for r_i, row in enumerate(parties_table):
            bg = LIGHT_HEX if r_i % 2 == 0 else WHITE_HEX
            for c_i, val in enumerate(row):
                cell = pt.rows[r_i+1].cells[c_i]
                set_cell_bg(cell, bg)
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(val))
                run.font.name = "Calibri"; run.font.size = Pt(9)
                if c_i == 0: run.bold = True; run.font.color.rgb = NAVY
        for col_w, col_i in zip([1.2, 3.3, 2.0], range(3)):
            for row in pt.rows:
                row.cells[col_i].width = Inches(col_w)
        doc.add_paragraph()

    # ── Status badge ──
    p_status = doc.add_paragraph()
    p_status.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_s = p_status.add_run(f"  {status}  ")
    r_s.font.name = "Calibri"; r_s.font.size = Pt(10); r_s.bold = True
    r_s.font.color.rgb = WHITE
    pPr2 = p_status._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear"); shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), BLUE_HEX)
    pPr2.append(shd2)
    p_status.paragraph_format.space_after = Pt(6)

    # ── Date ──
    p_date = doc.add_paragraph()
    r_d = p_date.add_run(f"Dated: {ac.get('lease_start', datetime.date.today().isoformat())}")
    r_d.font.name = "Calibri"; r_d.font.size = Pt(10); r_d.font.color.rgb = GRAY

    # ── Lessee logo (bottom right) ──
    doc.add_paragraph()
    doc.add_paragraph()
    p_ll = doc.add_paragraph()
    p_ll.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ll_buf = LESSEE_LOGOS[ac["lessee"]]
    ll_buf.seek(0)
    p_ll.add_run().add_picture(ll_buf, width=Inches(2.4))

    # ── Confidentiality notice ──
    add_horiz_rule(doc)
    p_conf = doc.add_paragraph()
    r_conf = p_conf.add_run(
        "This document is strictly confidential and is intended solely for the named parties. "
        "Any reproduction, disclosure, or distribution without the prior written consent of "
        f"{ac['lessor']} is strictly prohibited."
    )
    r_conf.font.name = "Calibri"; r_conf.font.size = Pt(8)
    r_conf.font.color.rgb = GRAY; r_conf.italic = True
    p_conf.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ─── Aircraft / Lessor / Lessee data ─────────────────────────────────────────

AIRCRAFT = [
    {"reg": "N12345", "type": "Boeing 737-800", "msn": "42301", "engines": "CFM56-7B27", "eng_esn": ["ESN-892341", "ESN-892342"], "year": 2018, "lessee": "SkyBridge Airlines Inc.", "lessor": "Aero Capital Finance Ltd.", "lease_start": "2024-01-15", "lease_end": "2030-01-14", "monthly_rent": 385000, "mr_rate_airframe": 145, "mr_rate_engine": 160, "mr_rate_llp": 210, "mr_rate_lg": 95, "mr_rate_apu": 55},
    {"reg": "N23456", "type": "Airbus A320-200", "msn": "8841", "engines": "CFM56-5B4/P", "eng_esn": ["ESN-771201", "ESN-771202"], "year": 2017, "lessee": "Meridian Express Corp.", "lessor": "Global Wing Finance BV", "lease_start": "2023-06-01", "lease_end": "2029-05-31", "monthly_rent": 360000, "mr_rate_airframe": 135, "mr_rate_engine": 155, "mr_rate_llp": 200, "mr_rate_lg": 90, "mr_rate_apu": 50},
    {"reg": "N34567", "type": "Boeing 737-900ER", "msn": "44501", "engines": "CFM56-7B27E", "eng_esn": ["ESN-910101", "ESN-910102"], "year": 2019, "lessee": "Apex Continental Airlines", "lessor": "Aero Capital Finance Ltd.", "lease_start": "2024-03-01", "lease_end": "2031-02-28", "monthly_rent": 420000, "mr_rate_airframe": 155, "mr_rate_engine": 170, "mr_rate_llp": 225, "mr_rate_lg": 100, "mr_rate_apu": 60},
    {"reg": "N45678", "type": "Airbus A321-200", "msn": "9102", "engines": "CFM56-5B3/P", "eng_esn": ["ESN-654321", "ESN-654322"], "year": 2020, "lessee": "Pacific Rim Air Travel Ltd.", "lessor": "Horizon Leasing Partners LP", "lease_start": "2024-07-01", "lease_end": "2031-06-30", "monthly_rent": 445000, "mr_rate_airframe": 160, "mr_rate_engine": 175, "mr_rate_llp": 230, "mr_rate_lg": 105, "mr_rate_apu": 65},
    {"reg": "N56789", "type": "Boeing 737 MAX 8", "msn": "45912", "engines": "CFM LEAP-1B27", "eng_esn": ["ESN-100201", "ESN-100202"], "year": 2022, "lessee": "Nordic Skies AS", "lessor": "Atlas Aircraft Trust", "lease_start": "2024-09-01", "lease_end": "2032-08-31", "monthly_rent": 510000, "mr_rate_airframe": 175, "mr_rate_engine": 195, "mr_rate_llp": 250, "mr_rate_lg": 115, "mr_rate_apu": 70},
    {"reg": "N67890", "type": "Airbus A319-100", "msn": "7203", "engines": "CFM56-5B5/P", "eng_esn": ["ESN-334455", "ESN-334456"], "year": 2016, "lessee": "Coastal Charter Services Ltd.", "lessor": "Global Wing Finance BV", "lease_start": "2022-11-01", "lease_end": "2028-10-31", "monthly_rent": 295000, "mr_rate_airframe": 120, "mr_rate_engine": 140, "mr_rate_llp": 185, "mr_rate_lg": 85, "mr_rate_apu": 45},
]

ATAS = [
    "ATA 05 – Time Limits/Maintenance Checks","ATA 06 – Dimensions and Areas",
    "ATA 07 – Lifting and Shoring","ATA 10 – Parking and Mooring",
    "ATA 11 – Placards and Markings","ATA 12 – Servicing",
    "ATA 21 – Air Conditioning","ATA 22 – Auto Flight",
    "ATA 23 – Communications","ATA 24 – Electrical Power",
    "ATA 25 – Equipment and Furnishings","ATA 26 – Fire Protection",
    "ATA 27 – Flight Controls","ATA 28 – Fuel",
    "ATA 29 – Hydraulic Power","ATA 30 – Ice and Rain Protection",
    "ATA 31 – Instruments","ATA 32 – Landing Gear",
    "ATA 33 – Lights","ATA 34 – Navigation",
    "ATA 35 – Oxygen","ATA 36 – Pneumatic",
    "ATA 38 – Water and Waste","ATA 45 – Central Maintenance System",
    "ATA 46 – Information Systems","ATA 49 – APU",
    "ATA 51 – Structures","ATA 52 – Doors",
    "ATA 53 – Fuselage","ATA 54 – Nacelles / Pylons",
    "ATA 55 – Stabilizers","ATA 56 – Windows",
    "ATA 57 – Wings","ATA 71 – Power Plant",
    "ATA 72 – Engine","ATA 73 – Engine Fuel and Control",
    "ATA 74 – Ignition","ATA 75 – Air",
    "ATA 76 – Engine Controls","ATA 80 – Starting",
]

LLP_MODULES = [
    {"module": "Fan", "parts": [
        {"pn": "1794M90P01", "desc": "Fan Disk", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1794M91P01", "desc": "Fan Blade (set of 24)", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1794M92P01", "desc": "Fan Case", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
    ]},
    {"module": "LPC (Booster)", "parts": [
        {"pn": "1795M10P01", "desc": "Booster Disk Stage 1", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1795M11P01", "desc": "Booster Disk Stage 2", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1795M12P01", "desc": "Booster Disk Stage 3", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1795M13P01", "desc": "Booster Spool", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
    ]},
    {"module": "HPC", "parts": [
        {"pn": "1796M01P01", "desc": "HPC Disk Stage 3", "life": 15000, "csn": 14822, "crsn": 14822, "remaining": 178},
        {"pn": "1796M02P01", "desc": "HPC Disk Stage 4", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1796M03P01", "desc": "HPC Disk Stage 5", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1796M04P01", "desc": "HPC Disk Stage 6", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1796M05P01", "desc": "HPC Disk Stage 7", "life": 25000, "csn": 14822, "crsn": 14822, "remaining": 10178},
        {"pn": "1796M06P01", "desc": "HPC Disk Stage 8", "life": 25000, "csn": 14822, "crsn": 14822, "remaining": 10178},
        {"pn": "1796M07P01", "desc": "HPC Disk Stage 9", "life": 25000, "csn": 14822, "crsn": 14822, "remaining": 10178},
        {"pn": "1796M08P01", "desc": "HPC Spool (3-4)", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1796M09P01", "desc": "HPC Spool (5-9)", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
    ]},
    {"module": "HPT", "parts": [
        {"pn": "1797M01P01", "desc": "HPT Disk Stage 1", "life": 15000, "csn": 14822, "crsn": 14822, "remaining": 178},
        {"pn": "1797M02P01", "desc": "HPT Disk Stage 2", "life": 20000, "csn": 14822, "crsn": 14822, "remaining": 5178},
        {"pn": "1797M03P01", "desc": "HPT Stage 1 Blade (set)", "life": 10000, "csn": 14822, "crsn": 2500, "remaining": 7500},
        {"pn": "1797M04P01", "desc": "HPT Stage 2 Blade (set)", "life": 10000, "csn": 14822, "crsn": 2500, "remaining": 7500},
    ]},
    {"module": "LPT", "parts": [
        {"pn": "1798M01P01", "desc": "LPT Disk Stage 4", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1798M02P01", "desc": "LPT Disk Stage 5", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1798M03P01", "desc": "LPT Disk Stage 6", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1798M04P01", "desc": "LPT Disk Stage 7", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
        {"pn": "1798M05P01", "desc": "LPT Shaft", "life": 30000, "csn": 14822, "crsn": 14822, "remaining": 15178},
    ]},
]


# ─── Document builders ────────────────────────────────────────────────────────

def build_lease_agreement(ac, idx):
    doc = Document()
    ref = f"ALA-{2024+idx:04d}-{idx:03d}"
    add_header_footer(doc, ref, "Aircraft Lease Agreement", ac["lessor"])
    add_cover_page(doc, "Aircraft Lease Agreement", ref, ac,
                   f"{ac['type']} — MSN {ac['msn']}",
                   [("Lessor", ac["lessor"], "Ireland"),
                    ("Lessee", ac["lessee"], "United States of America")],
                   status="EXECUTION COPY")

    heading(doc, "ARTICLE 1 — DEFINITIONS", 2); add_horiz_rule(doc)
    body(doc, "As used in this Agreement, the following terms shall have the meanings set forth below:")
    definitions = [
        ("Airframe", f"The {ac['type']} airframe bearing MSN {ac['msn']} and registration {ac['reg']}, together with all installed parts, appliances, avionics, and furnishings, but excluding the Engines."),
        ("Aircraft", "The Airframe together with the Engines, all as described in Schedule 1."),
        ("Base Rent", f"USD {ac['monthly_rent']:,} per calendar month, subject to annual CPI-U escalation."),
        ("C-Check Interval", "Nominally every 6 years / 36,000 FH / 5,000 FC, whichever occurs first, per the Lessee's approved Maintenance Programme."),
        ("Delivery Date", f"{ac['lease_start']}, being the date the Aircraft was delivered to Lessee."),
        ("Engine", f"Each {ac['engines']} turbofan bearing ESNs {ac['eng_esn'][0]} (Pos.1) and {ac['eng_esn'][1]} (Pos.2)."),
        ("Event of Default", "Any of the events described in Article 20."),
        ("Lease Period", f"{ac['lease_start']} to {ac['lease_end']}, unless earlier terminated."),
        ("LLP", "Life Limited Part — any component with a defined service life limit expressed in FH, FC, or calendar time."),
        ("Maintenance Reserves", "Supplemental rent per Article 8 and Schedule 3."),
        ("Return Conditions", "The physical redelivery standards described in Schedule 5."),
        ("Security Deposit", f"USD 1,155,000 (three months' Base Rent), payable before Delivery Date."),
        ("Total Loss", "Destruction, permanent loss, theft, or constructive total loss of the Aircraft."),
    ]
    add_table_styled(doc, ["Term", "Definition"], definitions, [1.8, 4.7])

    page_break(doc)
    heading(doc, "ARTICLE 2 — CONDITIONS PRECEDENT TO DELIVERY", 2); add_horiz_rule(doc)
    body(doc, "2.1 Lessor's obligation to deliver the Aircraft is conditional upon receipt of all of the following, in form satisfactory to Lessor, no later than five (5) Business Days before the Delivery Date:")
    cp_rows = [
        ("2.1(a)", "Evidence of incorporation and good standing of Lessee (certificate of good standing, within 30 days)"),
        ("2.1(b)", "Certified copies of all corporate authorisations (board resolutions, shareholder resolutions if required)"),
        ("2.1(c)", "Executed originals of each Transaction Document to which Lessee is a party"),
        ("2.1(d)", "Evidence of payment of the Security Deposit per Article 7"),
        ("2.1(e)", "Evidence of payment of first month's Base Rent per Article 6"),
        ("2.1(f)", "Certificates of insurance with required endorsements naming Lessor as additional insured and loss payee"),
        ("2.1(g)", "Legal opinion from Lessee's external counsel confirming due authorisation and enforceability"),
        ("2.1(h)", "Evidence of Aircraft registration at the FAA Civil Aircraft Register"),
        ("2.1(i)", "Copy of FAA-approved Maintenance Programme and current OpSpecs"),
        ("2.1(j)", "IOSA certification current within 24 months"),
        ("2.1(k)", "Executed Maintenance Reserves Account Agreement"),
    ]
    add_table_styled(doc, ["Ref", "Condition Precedent"], cp_rows, [0.6, 5.9])

    page_break(doc)
    heading(doc, "ARTICLE 3 — DELIVERY", 2); add_horiz_rule(doc)
    body(doc, f"3.1  Subject to satisfaction of the conditions in Article 2, Lessor shall make the Aircraft available for delivery at Dallas/Fort Worth International Airport (DFW) on or about {ac['lease_start']} (the 'Scheduled Delivery Date').")
    body(doc, "3.2  Lessee shall inspect the Aircraft within 48 hours of it being made available. Inspection shall be conducted by Lessee's technical representatives in the presence of Lessor's representative.")
    body(doc, "3.3  If Lessee identifies any non-conformity with the Delivery Condition, it shall issue a Discrepancy Notice. Lessor shall have 10 Business Days to remedy each non-conformity at its own cost.")
    body(doc, "3.4  Delivery is evidenced by execution of the Delivery Acceptance Certificate (Exhibit A). Execution constitutes Lessee's irrevocable acceptance, subject to express reservations.")
    body(doc, "3.5  Risk of loss and damage passes from Lessor to Lessee upon execution of the Delivery Acceptance Certificate.")
    body(doc, "3.6  Title to the Aircraft remains with Lessor at all times.")

    heading(doc, "ARTICLE 4 — LEASE PERIOD", 2); add_horiz_rule(doc)
    body(doc, f"4.1  The Lease Period commences on {ac['lease_start']} and continues until {ac['lease_end']}, unless earlier terminated.")
    body(doc, "4.2  Lessee has no option to extend or renew unless Lessor expressly agrees in writing no less than 12 months before the Expiry Date.")
    body(doc, "4.3  Holdover rate: 150% of the daily equivalent of Base Rent per day of holdover, without prejudice to other remedies.")

    page_break(doc)
    heading(doc, "ARTICLE 5 — PERMITTED USE", 2); add_horiz_rule(doc)
    body(doc, "5.1  Lessee shall operate the Aircraft solely in commercial air transport under Lessee's FAA Part 121 certificate and operations specifications.")
    body(doc, "5.2  No sub-lease, wet lease, dry lease, charter, or other use by a third party without Lessor's prior written consent.")
    body(doc, "5.3  No operation to/from/over sanctioned territories (OFAC, EU, UN).")
    body(doc, "5.4  No operation in areas of active military conflict or areas prohibited by Lessee's hull war insurers.")

    heading(doc, "ARTICLE 6 — RENT AND PAYMENT", 2); add_horiz_rule(doc)
    body(doc, f"6.1  Base Rent of USD {ac['monthly_rent']:,} per calendar month is payable in advance on the 1st of each month. Partial-month rent is prorated on a per-diem basis.")
    body(doc, "6.2  Base Rent escalates annually per CPI-U, with a floor of 0% and no ceiling.")
    body(doc, "6.3  All payments by wire transfer to:")
    add_table_styled(doc, ["Field", "Details"], [
        ("Bank", "J.P. Morgan Chase Bank N.A."),
        ("ABA", "021000021"),
        ("SWIFT", "CHASUS33"),
        ("Account Name", f"{ac['lessor']} — Aircraft Trust {ac['msn']}"),
        ("Account No.", f"1234567890{ac['msn'][-4:]}"),
        ("Reference", f"{ref} / {ac['reg']} / Month MMYYYY"),
    ], [1.8, 4.7])
    body(doc, "6.4  Late payment interest: SOFR + 4.00% p.a. on a 365-day basis from the due date to actual payment date.")
    body(doc, "6.5  All payments are without set-off, counterclaim, deduction, or withholding of any nature.")

    page_break(doc)
    heading(doc, "ARTICLE 7 — SECURITY DEPOSIT", 2); add_horiz_rule(doc)
    body(doc, f"7.1  Cash Security Deposit of USD 1,155,000 (three months' rent) payable before Delivery Date.")
    body(doc, "7.2  Held by Lessor in a segregated interest-bearing account; interest credited to Lessee if no Event of Default.")
    body(doc, "7.3  Lessor may apply any portion against unpaid amounts or losses; Lessee must replenish within 5 Business Days.")
    body(doc, "7.4  Returned within 30 days of redelivery, subject to no continuing defaults and all amounts paid in full.")
    body(doc, "7.5  May be replaced by a standby Letter of Credit from a Lessor-approved bank.")

    heading(doc, "ARTICLE 8 — MAINTENANCE RESERVES", 2); add_horiz_rule(doc)
    body(doc, "8.1  In addition to Base Rent, Lessee pays Maintenance Reserves as follows:")
    add_table_styled(doc, ["Reserve Component", "Rate", "Trigger", "Annual Escalation", "Cap (per event)"], [
        ("Airframe A/B Check", f"USD {ac['mr_rate_airframe']}/FH", "Per FH flown", "CPI-U", "USD 350,000"),
        ("Airframe C Check", f"USD {ac['mr_rate_airframe']*3}/FH", "Per FH flown", "CPI-U", "USD 4,500,000"),
        ("Engine LLP — No. 1", f"USD {ac['mr_rate_llp']}/FC", "Per FC on Eng 1", "CPI-U + 1%", "USD 2,800,000"),
        ("Engine LLP — No. 2", f"USD {ac['mr_rate_llp']}/FC", "Per FC on Eng 2", "CPI-U + 1%", "USD 2,800,000"),
        ("Engine Performance — No. 1", f"USD {ac['mr_rate_engine']}/FH", "Per FH on Eng 1", "CPI-U + 1%", "USD 5,200,000"),
        ("Engine Performance — No. 2", f"USD {ac['mr_rate_engine']}/FH", "Per FH on Eng 2", "CPI-U + 1%", "USD 5,200,000"),
        ("Landing Gear (all legs)", f"USD {ac['mr_rate_lg']}/FH", "Per FH flown", "CPI-U", "USD 900,000/leg"),
        ("APU", f"USD {ac['mr_rate_apu']}/APU-H", "Per APU hour", "CPI-U", "USD 350,000"),
    ], [1.8, 1.2, 1.2, 1.2, 1.1])
    body(doc, "8.2  Utilisation Report due within 10 Business Days after each month-end (form: Exhibit B).")
    body(doc, "8.3  MR Claim submission within 30 days of maintenance event completion, with full supporting documentation.")
    body(doc, "8.4  Lessor approves or disputes claims within 21 Business Days of receipt; eligible amounts paid within 5 Business Days of approval.")
    body(doc, "8.5  MR balances are non-refundable except as provided in Schedule 3.")

    page_break(doc)
    heading(doc, "ARTICLE 9 — AIRWORTHINESS AND MAINTENANCE", 2); add_horiz_rule(doc)
    body(doc, "9.1  Lessee shall maintain the Aircraft in airworthy condition in accordance with: (a) FAA regulations (14 C.F.R. Parts 21, 39, 43, 91, 121); (b) OEM manuals (AMM, IPC, AWM); (c) all applicable FAA and EASA Airworthiness Directives; (d) mandatory/alert Service Bulletins; (e) Lessee's approved Maintenance Programme; and (f) this Agreement.")
    body(doc, "9.2  All maintenance by FAA/EASA Part 145 approved MROs. Lessee to notify Lessor of MRO identity at least 10 Business Days before commencement.")
    body(doc, "9.3  Lessor has right of inspection on 5 Business Days' written notice (immediate for emergencies), without unreasonable interference with Lessee's operations.")

    heading(doc, "ARTICLE 10 — MODIFICATIONS", 2); add_horiz_rule(doc)
    body(doc, "10.1  No modifications without Lessor's prior written consent (not unreasonably withheld for AD-mandated modifications).")
    body(doc, "10.2  Approved modifications to be performed per approved data, not adversely affecting airworthiness or market value.")
    body(doc, "10.3  All modifications become Lessor's property. Lessee must restore to pre-modification configuration at redelivery unless Lessor directs otherwise.")

    heading(doc, "ARTICLE 11 — PARTS AND COMPONENTS", 2); add_horiz_rule(doc)
    body(doc, "11.1  Parts removed only for maintenance, repair, or replacement per this Agreement.")
    body(doc, "11.2  Replacement parts: like-for-like or better; FAA 8130-3 or EASA Form 1 release certificate required.")
    body(doc, "11.3  Title to replacement parts vests in Lessor upon installation, free and clear of all encumbrances.")

    page_break(doc)
    heading(doc, "ARTICLE 12 — TECHNICAL RECORDS", 2); add_horiz_rule(doc)
    body(doc, "12.1  All Technical Records maintained in English, in complete and current paper and electronic form per 14 C.F.R. Part 121 Subpart U.")
    body(doc, "12.2  Copies provided to Lessor on request.")
    body(doc, "12.3  All original Technical Records and electronic back-ups delivered to Lessor at redelivery.")

    heading(doc, "ARTICLE 13 — REGISTRATION AND TITLE", 2); add_horiz_rule(doc)
    body(doc, "13.1  Aircraft registered on the FAA Civil Aircraft Register throughout the Lease Period.")
    body(doc, "13.2  Registered in Lessor's name as owner and Lessee's name as operator.")
    body(doc, "13.3  No liens, mortgages, or charges to be registered against the Aircraft.")

    heading(doc, "ARTICLE 14 — LIENS", 2); add_horiz_rule(doc)
    body(doc, "14.1  Lessee shall not create or permit any lien on the Aircraft, its components, or Lessee's rights under this Agreement.")
    body(doc, "14.2  Any threatened lien must be discharged or bonded over within 5 Business Days of notification.")

    page_break(doc)
    heading(doc, "ARTICLE 15 — INSURANCE", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Coverage", "Minimum Limit", "Form / Terms", "Deductible", "Key Endorsements"], [
        ("Hull All Risks", "USD 52,000,000 Agreed Value", "AV1A + AVN52H", "USD 250,000", "Lessor loss payee; waiver of subrogation; no co-insurance"),
        ("Hull War and Allied Perils", "USD 52,000,000 Agreed Value", "AVN52H", "USD 500,000", "Confiscation/requisition; 7 days' notice"),
        ("Third-Party Liability", "USD 750,000,000 per occurrence", "AVN1C", "Nil", "Lessor additional insured; cross-liability; severability"),
        ("Passenger Liability", "USD 300,000 per seat (168 seats)", "Montreal Convention", "Nil", "IATA minimum compliant"),
        ("Baggage and Cargo", "USD 15,000,000 per occurrence", "Montreal Convention", "Nil", "—"),
        ("Workers' Comp / EL", "Statutory + USD 10M EL", "Standard endorsement", "Nil", "Waiver of subrogation"),
        ("Terrorism", "USD 52,000,000", "AVN52H", "USD 250,000", "7 days' cancellation notice"),
    ], [1.5, 1.3, 1.0, 0.9, 2.8])
    body(doc, "15.2  All insurers rated minimum A-/VII by A.M. Best.")
    body(doc, "15.3  Certificate of Insurance and endorsements to Lessor at least 5 Business Days before each policy period.")
    body(doc, "15.4  30 days' (7 days for war risk) cancellation notice to Lessor required under all policies.")

    page_break(doc)
    heading(doc, "ARTICLE 16 — INDEMNITIES", 2); add_horiz_rule(doc)
    body(doc, "16.1  General Indemnity. Lessee indemnifies each Indemnitee against all Claims arising out of: (a) manufacture, purchase, delivery, condition, use, maintenance, or redelivery of the Aircraft; (b) any breach by Lessee; (c) any act or omission of Lessee; (d) any third-party claim from Lessee's operations.")
    body(doc, "16.2  Tax Indemnity. Lessee to gross-up all payments against any withholding or deduction imposed by any Governmental Authority, other than taxes on Lessor's net income in its jurisdiction of incorporation.")

    heading(doc, "ARTICLE 17 — TOTAL LOSS", 2); add_horiz_rule(doc)
    body(doc, "17.1  Agreement terminates on the Total Loss Date. Lessee must immediately notify Lessor.")
    body(doc, "17.2  All insurance proceeds payable to Lessor. Surplus above Agreed Value paid to Lessee after discharge of all outstanding amounts.")
    body(doc, "17.3  Partial damage proceeds held by Lessor and released for repairs upon verification of airworthiness.")

    page_break(doc)
    heading(doc, "ARTICLE 18 — RETURN AND REDELIVERY", 2); add_horiz_rule(doc)
    body(doc, f"18.1  Aircraft redelivered to Lessor on {ac['lease_end']} (or earlier termination date) at a location to be agreed in the Continental United States or Western Europe.")
    body(doc, "18.2  Redelivery schedule agreed 180 days in advance, covering: (i) final maintenance check; (ii) technical inspection; (iii) redelivery flight.")
    body(doc, "18.3  Return Condition standards include: (a) valid FAA Airworthiness Certificate; (b) C-Check within 24 months; (c) each Engine PRSV within 36 months or <2,500 cycles since SV; (d) LLPs with >2,500 FC remaining; (e) Landing Gear >24 months/2,000 FC to next overhaul; (f) APU HSI within 24 months; (g) cabin in Delivery configuration or better; (h) exterior paint no primer visible.")
    body(doc, "18.4  Compensation in lieu payable per Schedule 5, Part II if Return Conditions are not met.")

    heading(doc, "ARTICLE 19 — REPRESENTATIONS AND WARRANTIES", 2); add_horiz_rule(doc)
    body(doc, "19.1  Each Party represents and warrants that: (a) duly incorporated and validly existing; (b) full authority to enter into and perform this Agreement; (c) Agreement constitutes its legal, valid, and binding obligation; (d) no consents outstanding; and (e) no conflict with applicable law.")

    page_break(doc)
    heading(doc, "ARTICLE 20 — EVENTS OF DEFAULT", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref", "Event", "Description", "Cure Period"], [
        ("20.1(a)", "Payment Default", "Failure to pay any amount on its due date", "5 Business Days from written notice"),
        ("20.1(b)", "Maintenance Default", "Failure to perform any maintenance obligation", "15 days (or longer if diligently pursuing remedy)"),
        ("20.1(c)", "Insurance Default", "Insurance lapses, cancelled, or fails to cover required risks", "3 Business Days"),
        ("20.1(d)", "Registration Default", "Airworthiness Certificate or registration suspended or revoked", "Immediate"),
        ("20.1(e)", "Insolvency", "General assignment for benefit of creditors, adjudicated bankrupt, receiver appointed", "None"),
        ("20.1(f)", "Cessation of Business", "Ceases commercial aviation or surrenders operating certificates", "None"),
        ("20.1(g)", "Cross-Default", "Default under any other agreement with Lessor or affiliate", "Per applicable grace period"),
        ("20.1(h)", "Misrepresentation", "Any representation proves materially false when made", "None"),
        ("20.1(i)", "Sanctions", "Lessee or operator becomes a Sanctioned Person", "None"),
        ("20.1(j)", "General Breach", "Material breach of any other obligation", "21 days from written notice"),
    ], [0.5, 1.4, 2.8, 1.8])
    body(doc, "20.2  Upon Event of Default, Lessor may: (a) terminate by written notice; (b) demand immediate redelivery; (c) repossess the Aircraft; (d) draw on Security Deposit/Letter of Credit; (e) claim all amounts due plus damages for unexpired Lease Period; and/or (f) exercise any other remedy available at law.")
    body(doc, "20.3  Lessee grants Lessor an irrevocable power of attorney, operative only upon Event of Default and termination, to repossess, deregister, and export the Aircraft.")

    page_break(doc)
    heading(doc, "ARTICLE 21 — ASSIGNMENT", 2); add_horiz_rule(doc)
    body(doc, "21.1  Lessor may assign all rights and interests to any financial institution as security for financing without Lessee's consent.")
    body(doc, "21.2  Lessee may not assign, transfer, or sub-lease any interest without Lessor's prior written consent (absolute discretion).")

    heading(doc, "ARTICLE 22 — GOVERNING LAW AND DISPUTE RESOLUTION", 2); add_horiz_rule(doc)
    body(doc, "22.1  New York law governs, without reference to conflict of laws principles.")
    body(doc, "22.2  Disputes resolved by ICC arbitration in New York (3 arbitrators, English language).")
    body(doc, "22.3  Emergency or interim relief from courts of competent jurisdiction is not precluded.")

    heading(doc, "ARTICLES 23–24 — NOTICES / MISCELLANEOUS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Party", "Address", "Email", "Attention"], [
        ("Lessor", f"{ac['lessor']}, 1 Grand Canal Square, Dublin 2, Ireland", "leases@lessor.ie", "Head of Portfolio Management"),
        ("Lessee", f"{ac['lessee']}, 123 Aviation Blvd, Dallas TX 75201", "leasing@lessee.com", "VP Finance & Leasing"),
    ], [1.2, 2.2, 1.8, 1.3])
    body(doc, "This Agreement constitutes the entire agreement and supersedes all prior negotiations. Amendments require a signed written instrument. Counterparts with electronic signatures are originals.")

    page_break(doc)
    heading(doc, "SCHEDULE 1 — AIRCRAFT DESCRIPTION", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Detail"], [
        ("Aircraft Type", ac["type"]),
        ("Manufacturer", "The Boeing Company" if "Boeing" in ac["type"] else "Airbus SAS"),
        ("MSN", ac["msn"]),
        ("Registration Mark", ac["reg"]),
        ("Year of Manufacture", str(ac["year"])),
        ("Engine Type", ac["engines"]),
        ("Engine Serial No. 1 (Position 1)", ac["eng_esn"][0]),
        ("Engine Serial No. 2 (Position 2)", ac["eng_esn"][1]),
        ("APU Model", "Honeywell 131-9(B)" if "737" in ac["type"] else "Pratt & Whitney APS3200"),
        ("APU Serial No.", f"APU-{ac['msn'][-4:]}01"),
        ("NLG Serial No.", f"NLG-{ac['msn'][-4:]}"),
        ("LH MLG Serial No.", f"LH-{ac['msn'][-4:]}"),
        ("RH MLG Serial No.", f"RH-{ac['msn'][-4:]}"),
        ("Interior Config.", "168 passengers (16 Business / 152 Economy)"),
        ("MTOW", "79,016 kg" if "737" in ac["type"] else "78,000 kg"),
        ("Operating Empty Weight", "41,413 kg" if "737" in ac["type"] else "42,600 kg"),
    ], [2.5, 4.0])

    heading(doc, "SCHEDULE 2 — TECHNICAL RECORDS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Record Type", "Description", "Format"], [
        ("Aircraft Logbook", "Continuous log of FH, FC, modifications, major maintenance", "Original + PDF scan"),
        ("Engine Logbooks (×2)", "Continuous log per engine FH, FC, TSO, shop visit history", "Original + PDF scan"),
        ("APU Logbook", "APU hours, cycles, shop visit history", "Original + PDF scan"),
        ("Airworthiness Directives", "Compliance record for all applicable ADs incl. repetitive items", "Electronic DB + paper"),
        ("Service Bulletin Status", "Master list of applicable SBs with incorporated/not status", "Electronic spreadsheet"),
        ("Component Status List", "All tracked components: P/N, S/N, TSO, TSI, time remaining", "Electronic CSV"),
        ("LLP Tracking Sheets", "Module-by-module LLP tracking per engine", "Electronic CSV"),
        ("C-Check Work Package", "Complete work card archive from last C-Check", "PDF, indexed"),
        ("Engine SV Reports", "Teardown and build-up reports from last PRSV", "PDF scans"),
        ("Modification / STC Records", "All STCs, FAA 337s, approved mod data", "Original + PDF"),
        ("Weight & Balance Report", "Current W&B issued by approved organisation", "Original + PDF"),
        ("Avionics Software Status", "Part numbers and versions of all LRU software", "Electronic"),
        ("Deferred Defect Log", "Open MEL/CDL items at delivery", "Electronic + paper"),
    ], [2.0, 3.2, 1.3])

    page_break(doc)
    heading(doc, "SCHEDULE 3 — MAINTENANCE RESERVE RATES", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component", "Rate", "Escalation", "Cap", "Eligible Cost Categories"], [
        ("Airframe A/B Check", f"USD {ac['mr_rate_airframe']}/FH", "CPI-U", "USD 350,000/event", "Labour, parts, NDT, consumables per AMM"),
        ("Airframe C Check", f"USD {ac['mr_rate_airframe']*3}/FH", "CPI-U", "USD 4,500,000/event", "As above plus structural repairs"),
        ("Engine LLP – No. 1", f"USD {ac['mr_rate_llp']}/FC", "CPI-U + 1%", "USD 2,800,000/module set", "OEM LLP parts with 8130-3/Form 1"),
        ("Engine LLP – No. 2", f"USD {ac['mr_rate_llp']}/FC", "CPI-U + 1%", "USD 2,800,000/module set", "OEM LLP parts with 8130-3/Form 1"),
        ("Engine Performance – No. 1", f"USD {ac['mr_rate_engine']}/FH", "CPI-U + 1%", "USD 5,200,000/SV", "PR workscope per PRSV tear-down findings"),
        ("Engine Performance – No. 2", f"USD {ac['mr_rate_engine']}/FH", "CPI-U + 1%", "USD 5,200,000/SV", "PR workscope per PRSV tear-down findings"),
        ("Main Landing Gear", f"USD {ac['mr_rate_lg']}/FH", "CPI-U", "USD 900,000/leg", "Overhaul to OEM spec; NDT; chrome plating"),
        ("Nose Landing Gear", f"USD {int(ac['mr_rate_lg']*0.6)}/FH", "CPI-U", "USD 400,000/overhaul", "Overhaul to OEM spec"),
        ("APU", f"USD {ac['mr_rate_apu']}/APU-H", "CPI-U", "USD 350,000/HSI", "HSI parts and labour"),
    ], [1.7, 1.0, 0.9, 1.3, 2.6])
    body(doc, "Ineligible costs: cosmetic improvements beyond Return Condition; operator preference modifications; FOD / bird strike damage (insured events); management fees; ferry costs beyond nearest capable MRO.")

    heading(doc, "SCHEDULE 4 — MINIMUM MAINTENANCE OBLIGATIONS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Check Type", "Interval", "Minimum at Redelivery"], [
        ("Transit / Pre-flight", "Per flight", "Per Lessee's approved checks"),
        ("Daily Check", "Each calendar day", "Per AMM Task 05-10-01"),
        ("A Check", "600 FH or 400 FC or 90 days (earliest)", "No open items; all ADs current"),
        ("B Check", "3,000 FH or 18 months (earliest)", "Full systems test; no MEL beyond Cat A"),
        ("C Check", "6,000 FH or 36 months or 5,000 FC (earliest)", "CPCP tasks; SRM repairs; NDT"),
        ("D Check", "24,000 FH or 12 years (earliest)", "Full strip and inspection per AMM"),
        ("Engine PRSV", "On-condition; max 8,000 FH from last SV", "Min EGT margin 30°C at delivery"),
        ("Engine LLP Replacement", "Per OEM life limits", "Min 2,500 FC remaining at redelivery"),
        ("Landing Gear Overhaul", "10 years or per manufacturer limits", "Min 24 months remaining at redelivery"),
        ("APU Hot Section Inspection", "4,000 APU-H or 24 months (earliest)", "Min 1,000 APU-H remaining at redelivery"),
    ], [1.8, 1.8, 2.9])

    page_break(doc)
    heading(doc, "SCHEDULE 5 — RETURN CONDITIONS AND COMPENSATION RATES", 2); add_horiz_rule(doc)
    body(doc, "Part I — Return Condition Standards", italic=False)
    add_table_styled(doc, ["Component", "Minimum at Redelivery", "Measurement"], [
        ("Airframe C Check", "Within preceding 24 months", "Work package sign-off date"),
        ("Structural Repairs", "No open deferred structural repairs; all CPCP current", "Compliance record"),
        ("Engine No. 1 EGT Margin", "Min 30°C (cruise, ISA)", "Borescope + performance run"),
        ("Engine No. 1 Shop Visit", "Within preceding 36 months", "SV completion records"),
        ("Engine No. 1 LLPs", "Min 2,500 FC remaining on all LLPs", "Engine records / disc sheets"),
        ("Engine No. 2 EGT Margin", "Min 30°C (cruise, ISA)", "Borescope + performance run"),
        ("Engine No. 2 Shop Visit", "Within preceding 36 months", "SV completion records"),
        ("Engine No. 2 LLPs", "Min 2,500 FC remaining on all LLPs", "Engine records / disc sheets"),
        ("APU", "HSI within preceding 24 months", "HSI work order and certificate"),
        ("Main Landing Gear", "Overhaul within 10 years; min 24 months/2,000 FC remaining", "Overhaul records"),
        ("Interior – Seats", "All serviceable; no unrepaired damage; covers clean/replaced", "Physical inspection"),
        ("Exterior Paint", "Full repaint if >25% oxidation or flaking; no primer visible", "Physical inspection"),
        ("MEL Items", "No open Cat A; Cat B/C within window", "MEL log"),
    ], [2.0, 2.5, 2.0])
    body(doc, "Part II — Compensation Rates (per shortfall unit)")
    add_table_styled(doc, ["Item", "Compensation Rate"], [
        ("C Check — each month short of 24-month minimum", "USD 62,000 / month"),
        ("Engine PRSV — each month short of 36-month minimum", "USD 45,000 / month / engine"),
        (f"Engine LLP — each FC short of 2,500 minimum (per LLP)", f"USD {ac['mr_rate_llp']} / FC / LLP"),
        ("Landing Gear — each month short of 24-month minimum", "USD 18,000 / month / leg"),
        ("APU — each month short of 24-month minimum", "USD 8,000 / month"),
        ("Exterior repaint required", "USD 185,000 lump sum"),
        ("Seat cover replacement", "USD 1,200 / seat"),
        ("Full carpet replacement", "USD 28,000 lump sum"),
        ("Open MEL item beyond window", "USD 5,000 / item / day until closed"),
    ], [3.5, 3.0])

    page_break(doc)
    sig_block(doc, [
        ("LESSOR", "James O'Sullivan", "Managing Director", ac["lease_start"]),
        ("LESSEE", "Catherine M. Rodriguez", "EVP Finance", ac["lease_start"]),
    ])

    path = os.path.join(OUT, "01-lease-agreements", f"LA-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: LA-{ac['reg']}-{ac['msn']}.docx")


def build_dcr(ac, idx):
    doc = Document()
    ref = f"DCR-{ac['reg']}-{ac['lease_start']}"
    add_header_footer(doc, ref, "Delivery Condition Report", ac["lessor"])
    add_cover_page(doc, "Delivery Condition Report", ref, ac,
                   f"Aircraft Delivery — {ac['reg']} / MSN {ac['msn']}",
                   [("Lessor", ac["lessor"], "Ireland"),
                    ("Lessee", ac["lessee"], "USA"),
                    ("Delivery Location", "Dallas/Fort Worth Int'l Airport (DFW)", "Texas, USA")],
                   status="SIGNED ORIGINAL")

    heading(doc, "1. DELIVERY SUMMARY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Detail"], [
        ("Aircraft Type", ac["type"]),
        ("MSN", ac["msn"]),
        ("Registration", ac["reg"]),
        ("Lessor", ac["lessor"]),
        ("Lessee", ac["lessee"]),
        ("Delivery Date", ac["lease_start"]),
        ("Delivery Location", "Dallas/Fort Worth Intl Airport (DFW), Texas, USA"),
        ("Inspector – Lessor", "Thomas Kavanaugh (AME Lic. 123456), TK Technical Services"),
        ("Inspector – Lessee", "Robert A. Fitch (IA Cert. 654321), Director of Technical"),
        ("Total FH at Delivery (TTSN)", "22,841 FH"),
        ("Total FC at Delivery (TCSN)", "16,203 FC"),
        ("FH since last C-Check", "4,102 FH"),
        ("Date of last C-Check", "2022-03-14"),
    ], [2.2, 4.3])

    page_break(doc)
    heading(doc, "2. ATA CHAPTER INSPECTION RESULTS", 2); add_horiz_rule(doc)
    body(doc, "Inspection findings by ATA chapter. Items with STATUS = 'Discrepancy' are detailed in Section 9.")
    statuses = ["Satisfactory"]*38 + ["Discrepancy", "Satisfactory"]
    notes_list = [
        "Per Mx Programme","Confirmed per ASM","All certified","All tied down","All legible, compliant",
        "All fluids serviced","Functional test OK","All AFCS tests pass","Slight wear in headliner lining",
        "All ELEC buses tested","Seat covers minor scuff","All detectors tested","All surfaces checked",
        "Sealants checked","Sys pressure normal","Boots/weep holes OK","EFIS/EICAS checked",
        "Tyre wear: 60% remain","Lighting functional","Nav DB current","Qty checked","All valves OK",
        "No leaks, potable water OK","CMS no active faults","IFE functional minor screen fault",
        "APU start-stop, load test","No disbonds observed","All latches tested","No visible cracks",
        "Cowl hinges OK","Elevator trim checked","All panels intact","No disbonds","ENG 1 start OK",
        "Blade FOD check negative","Fuel scheduling normal","Exciter test pass","Bleed valve tested",
        "Thrust lever friction HIGH — see Discrepancy DCR-001","Start cycle test pass"
    ]
    ata_rows = [(ATAS[i], statuses[i], notes_list[i]) for i in range(len(ATAS))]
    add_table_styled(doc, ["ATA Chapter", "Status", "Inspector Notes"], ata_rows, [2.2, 1.2, 3.1])

    page_break(doc)
    heading(doc, "3. ENGINE STATUS — ENGINE NO. 1 (ESN " + ac["eng_esn"][0] + ")", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Value"], [
        ("Engine Type", ac["engines"]),
        ("ESN", ac["eng_esn"][0]),
        ("Position", "No. 1 (Left / Wing Pylon 1)"),
        ("TTSN", "18,203 FH"),
        ("TCSN", "14,822 FC"),
        ("TSLSV (FH)", "4,102 FH / 2,918 FC"),
        ("Last SV Date", "2022-03-18"),
        ("Last SV Type", "Performance Restoration (PR) — CFM Approved"),
        ("Last SV Facility", "Delta TechOps, Atlanta GA — FAA OM1R074L / EASA 145.0049"),
        ("EGT Margin at Delivery", "38°C (minimum 30°C) — PASS"),
        ("Oil Consumption", "0.18 qt/hr (limit 0.50 qt/hr) — PASS"),
        ("Vibration N1/N2", "N1: 0.4 IPS / N2: 0.3 IPS — PASS"),
    ], [2.8, 3.7])

    heading(doc, "3.1 Engine No. 1 — LLP Status", 3)
    llp_rows = []
    for mod in LLP_MODULES:
        for p in mod["parts"]:
            pct = round(p["remaining"] / p["life"] * 100, 1)
            flag = "⚠ WARNING" if pct < 20 else "OK"
            llp_rows.append((mod["module"], p["pn"], p["desc"], str(p["life"]),
                             str(p["csn"]), str(p["crsn"]), str(p["remaining"]), f"{pct}%", flag))
    add_table_styled(doc, ["Module", "P/N", "Description", "Life", "TSN", "TSO", "Remaining", "% Left", "Flag"],
                     llp_rows, [0.8, 1.1, 1.5, 0.6, 0.6, 0.6, 0.7, 0.5, 0.8])

    page_break(doc)
    heading(doc, "4. ENGINE STATUS — ENGINE NO. 2 (ESN " + ac["eng_esn"][1] + ")", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Value"], [
        ("ESN", ac["eng_esn"][1]),
        ("Position", "No. 2 (Right / Wing Pylon 2)"),
        ("TTSN", "17,984 FH"),
        ("TCSN", "14,618 FC"),
        ("TSLSV", "3,841 FH / 2,714 FC"),
        ("Last SV Date", "2022-05-04"),
        ("Last SV Facility", "HAECO Xiamen — FAA MR7R078M / EASA 145.4170"),
        ("EGT Margin", "42°C — PASS"),
        ("Oil Consumption", "0.21 qt/hr — PASS"),
        ("Vibration", "N1: 0.5 IPS / N2: 0.4 IPS — PASS"),
    ], [2.8, 3.7])
    llp_rows2 = []
    for mod in LLP_MODULES:
        for p in mod["parts"]:
            csn2 = p["csn"] - 204
            rem2 = p["life"] - csn2
            pct2 = round(rem2 / p["life"] * 100, 1)
            flag2 = "⚠ WARNING" if pct2 < 20 else "OK"
            llp_rows2.append((mod["module"], p["pn"], p["desc"], str(p["life"]),
                              str(csn2), str(csn2), str(rem2), f"{pct2}%", flag2))
    add_table_styled(doc, ["Module","P/N","Description","Life","TSN","TSO","Remaining","% Left","Flag"],
                     llp_rows2, [0.8,1.1,1.5,0.6,0.6,0.6,0.7,0.5,0.8])

    page_break(doc)
    heading(doc, "5. APU STATUS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Value"], [
        ("APU Model", "Honeywell 131-9(B)"),
        ("APU S/N", f"APU-{ac['msn'][-4:]}01"),
        ("Total APU-H Since New", "12,441 APU-H"),
        ("Hours Since Last HSI", "2,108 APU-H"),
        ("Last HSI Date", "2022-09-15"),
        ("Last HSI Facility", "Honeywell Repair Station, Phoenix AZ — FAA OW1R089K"),
        ("Cold Start Test", "22 sec to idle (ambient 18°C) — PASS"),
        ("Bleed Air Load Test", "36 lb/min at sea level — PASS"),
        ("APU EGT at 100% load", "508°C (limit 685°C) — PASS"),
    ], [2.8, 3.7])

    heading(doc, "6. LANDING GEAR STATUS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Gear", "Serial No.", "Overhaul Date", "Facility", "FH Since OH", "FC Since OH", "Next OH Due"], [
        ("NLG", f"NLG-{ac['msn'][-4:]}", "2019-08-22", "AGES, Mexico City", "12,204 FH", "9,114 FC", "2029-08-22"),
        ("LH MLG", f"LH-{ac['msn'][-4:]}", "2020-02-14", "Messier Services, London Gatwick", "10,803 FH", "8,021 FC", "2030-02-14"),
        ("RH MLG", f"RH-{ac['msn'][-4:]}", "2020-02-14", "Messier Services, London Gatwick", "10,803 FH", "8,021 FC", "2030-02-14"),
    ], [0.9, 0.9, 1.0, 1.7, 0.9, 0.9, 1.3])

    page_break(doc)
    heading(doc, "7. CABIN INTERIOR SURVEY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Area", "Condition", "Defects Noted", "Action Required"], [
        ("Flight Deck", "Good", "None", "None"),
        ("Forward Galley (G1)", "Good", "Slight wear on trolley guides", "Monitor"),
        ("Business Class (16 seats)", "Good", "Seat 2A: armrest rattle", "Rectify before service"),
        ("Economy Class (132 seats)", "Fair", "Seat 14F: IFE screen blank; Seat 22B: pocket torn", "Rectify IFE; replace pocket"),
        ("Aft Galley (G2)", "Good", "None", "None"),
        ("Forward Lavatory", "Fair", "Waste flap hinge stiff", "Rectify"),
        ("Mid / Aft Lavatories (L2–L5)", "Good", "None", "None"),
        ("Overhead Bins", "Good", "Minor scuffing on door edges", "Acceptable"),
        ("Carpet – Forward", "Good", "No staining", "None"),
        ("Carpet – Economy", "Fair", "Minor wear rows 15–18", "Acceptable per Return Conditions"),
        ("Sidewall Panels", "Good", "2 panels with hairline cracks", "Monitor"),
        ("Emergency Equipment", "Serviceable", "All counts correct; no expiry within 12 months", "None"),
    ], [1.5, 0.9, 2.4, 1.7])

    heading(doc, "8. EXTERIOR CONDITION", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Zone", "Condition", "Noted Defects"], [
        ("Nose / Radome", "Good", "Paint chip 1.5 cm at lower radome — no structural damage"),
        ("Forward Fuselage (Sec 41/43)", "Good", "None"),
        ("Mid Fuselage (Sec 44/46)", "Good", "Door R2 outer seal slight extrusion — serviceable"),
        ("Aft Fuselage (Sec 46/48)", "Good", "None"),
        ("Tail Section / APU Bay", "Good", "None"),
        ("Vertical Stabiliser", "Good", "None"),
        ("Horizontal Stabilisers", "Good", "None"),
        ("Left Wing", "Good", "2 small paint scuffs (approx 3×5 cm total)"),
        ("Right Wing", "Good", "None"),
        ("Engine No. 1 Nacelle", "Good", "None"),
        ("Engine No. 2 Nacelle", "Good", "None"),
        ("Overall Paint", "Good", "~5% minor wear; no primer visible; no missing >2 cm"),
    ], [2.0, 1.0, 3.5])

    page_break(doc)
    heading(doc, "9. DISCREPANCY LOG", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref", "ATA", "Description", "Responsibility", "Resolution"], [
        ("DCR-001", "ATA 76", "Engine No.1 thrust lever friction high (8 lbf vs 6 lbf nominal)", "Lessor to rectify pre-delivery", "Adjust cable tension per AMM 76-11-00"),
        ("DCR-002", "ATA 25", "Seat 2A armrest rattle — Business Class", "Lessee within 30 days", "Replace armrest fitting; MEL Cat C"),
        ("DCR-003", "ATA 25", "Seat 14F IFE screen blank", "Lessee within 30 days", "Replace IFE LRU"),
        ("DCR-004", "ATA 25", "Seat 22B seatback pocket torn", "Lessee within 60 days", "Replace pocket fabric"),
        ("DCR-005", "ATA 38", "Forward lavatory waste flap hinge stiff", "Lessee within 30 days", "Lubricate/replace hinge per AMM"),
        ("DCR-006", "ATA 33", "Row 21 reading light inoperable — open MEL Cat C", "Lessee: open MEL acknowledged", "Replace at next A-Check"),
    ], [0.7, 0.7, 2.4, 1.3, 2.4])

    heading(doc, "10. OPEN MEL ITEMS AT DELIVERY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["MEL Ref", "ATA", "Description", "Cat", "Interval Expires"], [
        ("MEL-23-001", "ATA 23", "SELCAL code channel 4 inoperable", "B", "10 days from delivery"),
        ("MEL-25-007", "ATA 25", "Seat 14F IFE screen blank", "C", "30 days from delivery"),
        ("MEL-33-012", "ATA 33", "Row 21 reading light inoperable", "C", "120 days from delivery"),
    ], [1.0, 0.7, 2.8, 0.5, 1.5])

    page_break(doc)
    heading(doc, "11. TECHNICAL RECORDS COMPLETENESS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Record", "Present", "Complete", "Comments"], [
        ("Aircraft Logbook", "Yes", "Yes", "All entries legible and continuous"),
        ("Engine No. 1 Logbook", "Yes", "Yes", "Shop visit records attached"),
        ("Engine No. 2 Logbook", "Yes", "Yes", "Shop visit records attached"),
        ("APU Logbook", "Yes", "Yes", "All entries complete"),
        ("Component Status List", "Yes", "Yes", "Excel format; as at delivery date"),
        ("LLP Tracking Sheets (Eng 1 & 2)", "Yes", "Yes", "Disc sheets attached"),
        ("AD Compliance Record", "Yes", "Yes", "All ADs listed; repetitive due dates confirmed"),
        ("SB Status Record", "Yes", "Yes", "Alert SBs highlighted"),
        ("C-Check Work Package", "Yes", "Yes", "1,847 work cards archived"),
        ("Engine SV Reports", "Yes", "Yes", "Delta TechOps (Eng 1), HAECO Xiamen (Eng 2)"),
        ("Weight & Balance Report", "Yes", "Yes", "Current; dated within 12 months"),
        ("STCs / FAA 337s", "Yes", "Yes", "IFE STC SA11801SW, SATCOM STC SA10912SW"),
    ], [2.2, 0.7, 0.8, 2.8])

    sig_block(doc, [
        ("LESSOR REPRESENTATIVE", "Thomas Kavanaugh", "Technical Inspector", ac["lease_start"]),
        ("LESSEE REPRESENTATIVE", "Robert A. Fitch", "Director of Technical", ac["lease_start"]),
    ])

    path = os.path.join(OUT, "02-delivery-condition-reports", f"DCR-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: DCR-{ac['reg']}-{ac['msn']}.docx")


def build_mr_claim(ac, idx):
    doc = Document()
    ref = f"MRC-{ac['reg']}-2025-{idx:02d}"
    add_header_footer(doc, ref, "Maintenance Reserve Claim", ac["lessor"])
    add_cover_page(doc, "Maintenance Reserve Claim", ref, ac,
                   f"Engine PRSV — ESN {ac['eng_esn'][0]}",
                   [("Claimant (Lessee)", ac["lessee"], "USA"),
                    ("Reserve Account Holder", ac["lessor"], "Ireland"),
                    ("MRO Facility", "Lufthansa Technik AG", "Hamburg, Germany")],
                   status="CLAIM SUBMISSION")

    heading(doc, "1. CLAIM SUMMARY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Detail"], [
        ("Claim Reference", ref),
        ("Aircraft", f"{ac['type']} / {ac['reg']} / MSN {ac['msn']}"),
        ("Engine", f"{ac['engines']} / ESN {ac['eng_esn'][0]} (Position 1)"),
        ("Event Type", "Performance Restoration Shop Visit (PRSV)"),
        ("MRO Facility", "Lufthansa Technik AG, Hamburg — EASA 145.0001 / FAA OM3R164L"),
        ("Induction Date", "2025-01-08"),
        ("Completion Date", "2025-03-22"),
        ("Aircraft back in service", "2025-03-28"),
        ("Claim Submission Date", "2025-04-10"),
        ("Total Claim Amount", "USD 1,117,570"),
        ("MR Account Balance (Performance sub-account)", "USD 824,000"),
        ("MR Account Balance (LLP sub-account)", "USD 416,000"),
        ("Net Lessee Shortfall", "USD 0 (accounts cover full claim)"),
    ], [2.5, 4.0])

    page_break(doc)
    heading(doc, "2. REASON FOR SHOP VISIT", 2); add_horiz_rule(doc)
    body(doc, "The engine was inducted for a Performance Restoration shop visit following sustained EGT margin deterioration and a high-pressure compressor stall event on 2024-12-18. EGT margin had degraded to 12°C against a monitoring threshold of 20°C. An on-wing borescope conducted on 2024-12-22 revealed:")
    add_table_styled(doc, ["Ref", "ATA", "Finding", "Severity"], [
        ("BSI-001", "72-30", "HPT Stage 1 blades: 4 blades with tip erosion >AMM limits (>0.15 in)", "MANDATORY REPAIR"),
        ("BSI-002", "72-30", "HPT Stage 2 blades: 2 blades with leading edge nicks within limits — monitor", "MONITOR"),
        ("BSI-003", "72-40", "LPT Stage 4 Nozzle: 3 segments with oxidation burns exceeding AMM limits", "MANDATORY REPAIR"),
        ("BSI-004", "72-20", "HPC stages 5–7: performance coating loss contributing to EGT deterioration", "PRICED WORKSCOPE"),
        ("BSI-005", "72-10", "Fan blades: all within AMM limits; no damage", "SERVICEABLE"),
        ("BSI-006", "72-50", "No.1 and No.2 bearing compartment seals within limits", "SERVICEABLE"),
    ], [0.7, 0.8, 3.5, 1.5])

    page_break(doc)
    heading(doc, "3. AGREED WORKSCOPE", 2); add_horiz_rule(doc)
    workscope = [
        ("WS-01","72-00","Complete engine disassembly; module split into 5 modules","N/A","Done"),
        ("WS-02","72-10","Fan module: clean, inspect; restore tip clearances; replace 4 damaged fan blades","4 blades","Done"),
        ("WS-03","72-20","LPC booster: disassembly; clean; inspect; replace full blade set Stage 2","Full set Stg 2","Done"),
        ("WS-04","72-20","HPC: disassembly; SWET process on discs; replace HVOF coating Stages 5–7","Coating renewed","Done"),
        ("WS-05","72-20","HPC Disc Stage 3: replace (CSN at 14,990 — within 10 of limit per LLP sheet)","1 disc","Done"),
        ("WS-06","72-30","Combustion case: clean; inspect; 2 segments repaired per SRM 72-41-00","2 repairs","Done"),
        ("WS-07","72-30","HPT Stage 1 blades: full replacement set — all new, RFN traceable","Full set (80)","Done"),
        ("WS-08","72-30","HPT Stage 2 blades: replace 2 blades exceeding AMM limits","2 blades","Done"),
        ("WS-09","72-30","HPT Stage 1 Nozzle Guide Vanes: full ring replacement","Full ring","Done"),
        ("WS-10","72-30","HPT Stage 2 Nozzle Guide Vanes: full ring replacement","Full ring","Done"),
        ("WS-11","72-40","LPT Stage 4 Nozzle: 6 segments repaired per SRM; remaining 3 replaced","3 segments","Done"),
        ("WS-12","72-40","LPT blades Stages 4–7: clean, inspect; replace 8 blades (FOD/erosion)","8 blades","Done"),
        ("WS-13","72-50","No. 1 Bearing: replace oil seals and carbon seals; bearing inspected (no replacement)","Seals replaced","Done"),
        ("WS-14","72-50","No. 3 Bearing: replace (at CSN limit per LLP sheet)","1 bearing","Done"),
        ("WS-15","73-00","Fuel nozzles: replace all 20 (atomisation degraded)","20 nozzles","Done"),
        ("WS-16","74-00","Igniter plugs: replace both (hours on limit)","2 plugs","Done"),
        ("WS-17","72-00","Module reassembly; cold and hot section balance","N/A","Done"),
        ("WS-18","72-00","Performance test cell — 3 runs; EGT margin confirmed at 62°C — PASS","N/A","PASS"),
        ("WS-19","72-00","EASA Form 1 / FAA 8130-3 release; QA acceptance","N/A","Done"),
    ]
    add_table_styled(doc, ["Task","ATA","Work Description","Qty","Status"], workscope, [0.5,0.7,4.1,1.0,0.8])

    page_break(doc)
    heading(doc, "4. PARTS LIST", 2); add_horiz_rule(doc)
    parts = [
        ("1794M91P01","Fan Blade set (4 replaced)","4","NEW OEM","CFM/BOR2024","USD 42,800"),
        ("1795M11P01","LPC Blade set Stage 2 (28 blades)","28","NEW OEM","LHT/BOR2025","USD 28,600"),
        ("1796M01P01","HPC Disc Stage 3 (LLP at limit)","1","NEW OEM","CFM/LLP2025","USD 148,500"),
        ("1797M03P01","HPT Stage 1 Blade set (80 blades)","80","NEW OEM","LHT/BOR2025","USD 312,400"),
        ("1797M04P01","HPT Stage 2 Blade (2 blades)","2","NEW OEM","LHT/BOR2025","USD 18,600"),
        ("HPT-NGV-S1","HPT Stg 1 Nozzle Guide Vane Ring","1","NEW OEM","CFM/NGV2025","USD 89,200"),
        ("HPT-NGV-S2","HPT Stg 2 Nozzle Guide Vane Ring","1","NEW OEM","CFM/NGV2025","USD 72,100"),
        ("LPT-NOZ-S4","LPT Stage 4 Nozzle (3 segments)","3","NEW OEM","LHT/BOR2025","USD 34,500"),
        ("LPT-BLD-MIX","LPT Blades Stages 4–7 (8 blades)","8","NEW OEM","LHT/BOR2025","USD 22,800"),
        ("BRG-NO3-XX","No. 3 Bearing (LLP at limit)","1","NEW OEM","CFM/BRG2025","USD 8,400"),
        ("SEL-CRBNS-1","Carbon Seals No.1 Bearing Compartment","3","NEW OEM","LHT/SL2025","USD 4,200"),
        ("FNZ-20PLCE","Fuel Nozzles (set of 20)","20","NEW OEM","CFM/FN2025","USD 41,600"),
        ("IGN-PLUG-SET","Igniter Plug Set (2×)","2","NEW OEM","LHT/IG2025","USD 1,800"),
        ("HVOF-COAT-567","HVOF Coating HPC Stages 5–7","3 stages","PROCESS","LHT/HVOF2025","USD 36,800"),
        ("MISC-HARDWARE","Hardware (nuts, bolts, seals, gaskets)","Lot","NEW OEM","LHT/MISC2025","USD 12,400"),
    ]
    add_table_styled(doc, ["Part No.","Description","Qty","Cond.","Cert. Ref.","Cost"], parts, [1.1,2.3,0.4,0.8,1.2,1.1])

    heading(doc, "5. LABOUR BREAKDOWN", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Activity","Hours","Rate (USD/hr)","Total (USD)"], [
        ("Engine Disassembly & Assembly — Lead Mechanic","284","USD 145","USD 41,180"),
        ("Module Build Fan/LPC/HPC — Mechanic","312","USD 132","USD 41,184"),
        ("Hot Section Build HPT/LPT — Mechanic","298","USD 132","USD 39,336"),
        ("NDT (FPI/UT/Eddy Current)","96","USD 158","USD 15,168"),
        ("QA Inspection","48","USD 165","USD 7,920"),
        ("Test Cell Operator (3 runs)","36","USD 145","USD 5,220"),
        ("Documentation / Certification","24","USD 120","USD 2,880"),
        ("Overhead (15%)","—","—","USD 23,082"),
        ("TOTAL LABOUR","1,098 hrs","—","USD 175,970"),
    ], [2.5,1.0,1.3,1.7])

    page_break(doc)
    heading(doc, "6. COST SUMMARY AND ELIGIBILITY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Category","Amount","Eligible Account","Eligible?"], [
        ("OEM Parts — HPT/NGV/Fan/Fuel","USD 876,500","Engine Performance","YES"),
        ("HPC Stage 3 LLP replacement","USD 148,500","Engine LLP","YES"),
        ("No. 3 Bearing LLP replacement","USD 8,400","Engine LLP","YES"),
        ("Labour (1,098 hrs)","USD 175,970","Engine Performance","YES"),
        ("Test Cell (3 runs)","USD 38,500","Engine Performance","YES"),
        ("Transport (ferry + delivery)","USD 14,200","Engine Performance","YES"),
        ("Tooling","USD 12,400","Engine Performance","YES"),
        ("TOTAL CLAIM","USD 1,117,570","—","—"),
        ("Performance sub-account balance","USD 824,000","—","—"),
        ("LLP sub-account balance","USD 416,000","—","—"),
        ("Lessee shortfall","USD 0","—","COVERED"),
    ], [2.0,1.2,1.5,1.8])

    heading(doc, "7. SUPPORTING DOCUMENTATION ATTACHED", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref","Document","Pages"], [
        ("SUPP-01","LHT Shop Visit Work Order WO-LHT-2025-04812","48"),
        ("SUPP-02","Teardown Inspection Report (TIR)","32"),
        ("SUPP-03","Build-Up Inspection Report (BIR)","28"),
        ("SUPP-04","Performance Test Cell Run Report — Runs 1, 2, 3","18"),
        ("SUPP-05","EASA Form 1 / FAA 8130-3 for all replaced parts","31"),
        ("SUPP-06","Parts traceability back-to-birth documentation","22"),
        ("SUPP-07","Updated LLP tracking sheet (post-PRSV)","4"),
        ("SUPP-08","Updated Engine Logbook pages","8"),
        ("SUPP-09","Final Invoice — Lufthansa Technik AG","6"),
        ("SUPP-10","Photo documentation (188 photos)","38"),
        ("SUPP-11","Pre-PRSV borescope report (Delta TechOps)","14"),
    ], [0.8,4.5,0.7])

    sig_block(doc, [
        ("CLAIMANT (LESSEE)", "Sarah J. Watkins", "VP Technical Operations", "2025-04-10"),
        ("LESSOR (APPROVAL)", "Patrick Dunne", "Portfolio Technical Manager", "___________"),
    ])

    path = os.path.join(OUT, "03-maintenance-reserve-claims", f"MRC-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: MRC-{ac['reg']}-{ac['msn']}.docx")


def build_rcr(ac, idx):
    doc = Document()
    ref = f"RCR-{ac['reg']}-{ac['lease_end'][:7]}"
    add_header_footer(doc, ref, "Return Condition Report", ac["lessor"])
    add_cover_page(doc, "Return Condition Report", ref, ac,
                   f"Aircraft Return — Lease Expiry {ac['lease_end']}",
                   [("Lessor", ac["lessor"], "Ireland"),
                    ("Lessee", ac["lessee"], "USA"),
                    ("Return Location", "Miami Intl Airport (MIA)", "Florida, USA")],
                   status="FINAL — SIGNED")

    heading(doc, "1. RETURN OVERVIEW", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Detail"], [
        ("Aircraft", f"{ac['type']} / {ac['reg']} / MSN {ac['msn']}"),
        ("Lease Expiry", ac["lease_end"]),
        ("Return Date", ac["lease_end"]),
        ("Return Location", "Miami International Airport (MIA), FL, USA"),
        ("Lessor Inspector", "Declan Murphy (AME), Murphy Aviation Surveys Ltd"),
        ("Lessee Inspector", "James Parrish (A&P/IA), VP Engineering"),
        ("Independent Inspector", "AeroDynamic Advisory LLC"),
        ("Total FH Accumulated (Lease Period)", "15,363 FH"),
        ("Total FC Accumulated (Lease Period)", "11,611 FC"),
    ], [2.5, 4.0])

    page_break(doc)
    heading(doc, "2. AIRFRAME STATUS AT RETURN", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "At Delivery", "At Return", "Required", "Status"], [
        ("Total FH", "22,841", "38,204", "N/A", "INFO"),
        ("Total FC", "16,203", "27,814", "N/A", "INFO"),
        ("FH since C-Check", "4,102", "2,218", "Max 6,000 FH / 24 months", "PASS"),
        ("Date of last C-Check", "2022-03-14", "2029-11-08", "Within 24 months", "PASS"),
        ("Open AD items", "0", "0", "0 open", "PASS"),
        ("Open structural repairs", "0", "2", "0 open", "FAIL — compensation due"),
        ("Open MEL items", "3", "5", "Max Cat C in-window", "FAIL — 2 items overdue"),
    ], [1.8, 0.9, 0.9, 1.4, 0.8])

    heading(doc, "3. ENGINE STATUS AT RETURN", 2); add_horiz_rule(doc)
    for eng_label, esn, tcsn in [("No. 1", ac["eng_esn"][0], 27814), ("No. 2", ac["eng_esn"][1], 27610)]:
        heading(doc, f"3.{eng_label[3]}. Engine {eng_label} — ESN {esn}", 3)
        add_table_styled(doc, ["Parameter","Value","Required","Status"], [
            ("Date last PRSV", "2027-08-15", "Within 36 months", "PASS"),
            ("Months since last PRSV", "29 months", "Max 36 months", "PASS"),
            ("EGT Margin at Return", "22°C", "Min 30°C", "FAIL — 8°C shortfall"),
            ("Oil Consumption", "0.28 qt/hr", "Max 0.50 qt/hr", "PASS"),
            ("Vibration N1/N2", "N1: 0.6 / N2: 0.5 IPS", "Max 0.7 IPS each", "PASS"),
        ], [2.0, 1.5, 1.5, 0.8])

    page_break(doc)
    heading(doc, "4. ENGINE LLP STATUS AT RETURN", 2); add_horiz_rule(doc)
    body(doc, "LLPs below 2,500 FC remaining trigger compensation at USD " + str(ac["mr_rate_llp"]) + "/FC per LLP per lease Schedule 5.")
    llp_return = []
    for mod in LLP_MODULES:
        for p in mod["parts"]:
            csn_r = p["csn"] + 11611
            rem_r = max(0, p["life"] - csn_r)
            shortfall = max(0, 2500 - rem_r)
            comp = shortfall * ac["mr_rate_llp"]
            llp_return.append((
                mod["module"], p["pn"], p["desc"],
                str(p["life"]), str(csn_r), str(rem_r),
                str(shortfall), f"USD {comp:,}" if shortfall > 0 else "—",
                "FAIL" if rem_r < 2500 else "PASS"
            ))
    add_table_styled(doc, ["Module","P/N","Desc","Life","CSN","Remaining","Shortfall FC","Compensation","Status"],
                     llp_return, [0.7,1.0,1.5,0.6,0.7,0.7,0.7,0.9,0.5])

    page_break(doc)
    heading(doc, "5. COMPENSATION SUMMARY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Item","Calculation","Amount (USD)"], [
        ("Engine No.1 EGT shortfall (8°C)", "8°C × USD 45,000/month equivalent × 9 months = USD 405,000 per Sched. 5", "USD 405,000"),
        ("Engine No.2 EGT", "Meets minimum — no compensation due", "USD 0"),
        ("LLP shortfall (all modules combined)", f"Per LLP table above at USD {ac['mr_rate_llp']}/FC", "TBC per final LLP calc"),
        ("Open structural repairs (2)", "2 items × USD 5,000/day × 45 days estimated", "USD 450,000"),
        ("MEL overdue items (2)", "2 items × USD 5,000/day × 15 days estimated", "USD 150,000"),
        ("NLG overhaul overdue (4 months)", "4 months × USD 18,000/month", "USD 72,000"),
        ("Economy carpet replacement", "Lump sum per Schedule 5 Part II", "USD 28,000"),
        ("2 Business Class unserviceable seats", "Lump sum per Schedule 5 Part II", "USD 15,000"),
        ("IFE screens non-functional (14)", "14 × USD 3,000", "USD 42,000"),
        ("TOTAL COMPENSATION DUE FROM LESSEE", "", "USD 1,162,000+"),
    ], [2.5, 2.5, 1.5])

    heading(doc, "6. APU / LANDING GEAR STATUS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component","Value","Required","Status"], [
        ("APU Hours since HSI","1,802 APU-H","Max 4,000 / 24 months","PASS"),
        ("Last APU HSI Date","2028-04-10","Within 24 months (16 months ago)","PASS"),
        ("APU Cold Start","26 sec to idle","Max 30 sec","PASS"),
        ("NLG — months since overhaul","124 months","Max 120 months","FAIL — 4 months overdue"),
        ("LH MLG — months since overhaul","118 months","Max 120 months","PASS — 2 months remaining"),
        ("RH MLG — months since overhaul","118 months","Max 120 months","PASS — 2 months remaining"),
    ], [2.2, 1.5, 1.5, 1.6])

    sig_block(doc, [
        ("LESSOR INSPECTOR", "Declan Murphy", "Chief Inspector", ac["lease_end"]),
        ("LESSEE INSPECTOR", "James Parrish", "VP Engineering", ac["lease_end"]),
        ("INDEPENDENT INSPECTOR", "Dr. Clara Santos", "AeroDynamic Advisory", ac["lease_end"]),
    ])

    path = os.path.join(OUT, "04-return-condition-reports", f"RCR-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: RCR-{ac['reg']}-{ac['msn']}.docx")


def build_amendment(ac, idx):
    doc = Document()
    ref = f"AMEND-ALA-{2024+idx:04d}-{idx:03d}-No.{idx%3+1}"
    add_header_footer(doc, ref, "Lease Amendment Agreement", ac["lessor"])
    add_cover_page(doc, "Lease Amendment Agreement", ref, ac,
                   f"Amendment No. {idx%3+1} to Lease {ref[:20]}",
                   [("Lessor", ac["lessor"], "Ireland"),
                    ("Lessee", ac["lessee"], "USA")],
                   status="EXECUTED AMENDMENT")

    body(doc, f"THIS LEASE AMENDMENT AGREEMENT ('Amendment') is entered into as of 2025-06-15 and amends the Aircraft Lease Agreement dated {ac['lease_start']} ('Original Lease') between {ac['lessor']} ('Lessor') and {ac['lessee']} ('Lessee') for {ac['type']} MSN {ac['msn']}.")
    body(doc, "In the event of conflict between this Amendment and the Original Lease, this Amendment prevails. All defined terms have the meanings given in the Original Lease unless otherwise defined herein.")

    heading(doc, "ARTICLE 1 — AMENDMENT TO BASE RENT", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Clause","Original Text","Amended Text","Effective Date"], [
        ("Art. 6.1 — Base Rent", f"USD {ac['monthly_rent']:,}/month", f"USD {int(ac['monthly_rent']*1.032):,}/month (CPI-U 3.2% adjustment)", "1 July 2025"),
        ("Art. 6.2 — Escalation", "Annual CPI-U adjustment", "CPI-U with floor 2.0%, ceiling 5.0% p.a.", "1 July 2025"),
        ("Art. 6.4 — Late Interest", "SOFR + 4.00%", "SOFR + 3.50% (reduced — clean payment history)", "1 July 2025"),
    ], [1.2,2.2,2.2,0.9])
    body(doc, "Security Deposit adjusted to USD " + f"{int(ac['monthly_rent']*1.032*3):,}" + " within 30 days of effective date.")

    heading(doc, "ARTICLE 2 — AMENDMENT TO MAINTENANCE RESERVE RATES", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component","Original Rate","Amended Rate","Reason"], [
        ("Engine Performance – No.1", f"USD {ac['mr_rate_engine']}/FH", f"USD {int(ac['mr_rate_engine']*1.05)}/FH", "Annual CPI-U + 1%"),
        ("Engine Performance – No.2", f"USD {ac['mr_rate_engine']}/FH", f"USD {int(ac['mr_rate_engine']*1.05)}/FH", "Annual CPI-U + 1%"),
        ("Engine LLP – No.1", f"USD {ac['mr_rate_llp']}/FC", f"USD {int(ac['mr_rate_llp']*1.05)}/FC", "Annual CPI-U + 1%"),
        ("Engine LLP – No.2", f"USD {ac['mr_rate_llp']}/FC", f"USD {int(ac['mr_rate_llp']*1.05)}/FC", "Annual CPI-U + 1%"),
        ("Airframe C Check", f"USD {ac['mr_rate_airframe']*3}/FH", f"USD {int(ac['mr_rate_airframe']*3*1.032)}/FH", "CPI-U"),
        ("APU", f"USD {ac['mr_rate_apu']}/APU-H", f"USD {int(ac['mr_rate_apu']*1.032)}/APU-H", "CPI-U"),
    ], [1.8,1.2,1.2,2.3])

    page_break(doc)
    heading(doc, "ARTICLE 3 — AMENDMENT TO PERMITTED USE", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Activity","Original Position","Amended Position","Conditions"], [
        ("Wet Lease / ACMI","Required Lessor consent for each wet lease","Permitted without consent for up to 45 days / 12-month period","Lessee remains PIC operator; insurance covers ACMI; written notice within 5 BD"),
        ("Code-Share","Not addressed","Permitted with IOSA-registered carriers","Lessee remains operating carrier; no operational control transfer"),
        ("Charter","Permitted within OpSpecs","Confirmed permitted; ambiguity removed","Within existing OpSpecs; no new sub-certificate"),
    ], [1.2,1.5,1.7,2.1])

    heading(doc, "ARTICLE 4 — AMENDMENT TO RETURN CONDITIONS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component","Original Standard","Amended Standard","Rationale"], [
        ("Engine EGT Margin","Min 30°C at return","Min 25°C; compensation USD 45,000/month per degree below 30°C","Market practice; avoids unnecessary SV triggers"),
        ("Engine SV Recency","Within 36 months","Within 40 months","Greater scheduling flexibility"),
        ("Economy Carpet","Max 25% wear","Max 35% wear (no staining or odour)","Revised per industry benchmarking"),
        ("Engine LLP minimum","2,500 FC","2,000 FC; compensation at rate per cycle for shortfall below 2,500 FC","Market flexibility with compensation mechanism"),
    ], [1.5,1.3,1.6,2.1])

    heading(doc, "ARTICLE 5 — ADDITIONAL PROVISIONS", 2); add_horiz_rule(doc)
    body(doc, "5.1  Wi-Fi STC Pre-Approval: Lessor grants pre-approval for Inmarsat GX Aviation or equivalent Wi-Fi STC, subject to: (a) 30 days' advance notice with STC number and EO; (b) installation by FAA-approved repair station; (c) no adverse effect on any existing STC or airworthiness approval.")
    body(doc, "5.2  Engine Substitution: Temporary Replacement Engine permitted without consent if: (a) same type and model as removed Engine; (b) removed Engine inducted within 90 days; (c) Lessor notified immediately upon installation.")

    heading(doc, "ARTICLE 6 — RATIFICATION AND GENERAL", 2); add_horiz_rule(doc)
    body(doc, "All terms of the Original Lease not expressly modified by this Amendment remain in full force and effect and are hereby ratified. From the date hereof, all references to 'the Lease' shall mean the Original Lease as amended. This Amendment is governed by New York law.")

    sig_block(doc, [
        ("LESSOR", "Siobhan O'Keeffe", "VP Portfolio Transactions", "2025-06-15"),
        ("LESSEE", "Marcus T. Webb", "Chief Financial Officer", "2025-06-15"),
    ])

    path = os.path.join(OUT, "05-lease-amendments", f"AMEND-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: AMEND-{ac['reg']}-{ac['msn']}.docx")


def build_loi(ac, idx):
    doc = Document()
    ref = f"LOI-{ac['reg'][:2]}{idx:02d}-2026"
    add_header_footer(doc, ref, "Letter of Intent", ac["lessor"])
    add_cover_page(doc, "Letter of Intent", ref, ac,
                   f"Proposed Lease — {ac['type']} MSN {ac['msn']}",
                   [("Prospective Lessor", ac["lessor"], "Ireland"),
                    ("Prospective Lessee", ac["lessee"], "USA")],
                   status="NON-BINDING — SUBJECT TO ALA EXECUTION")

    body(doc, f"This Letter of Intent ('LOI') is submitted by {ac['lessee']} ('Prospective Lessee') to {ac['lessor']} ('Prospective Lessor') and sets out the non-binding commercial terms for the proposed lease of the Aircraft described below. This LOI does not constitute a binding obligation unless a formal Aircraft Lease Agreement ('ALA') is duly executed by both Parties.")

    heading(doc, "1. AIRCRAFT DESCRIPTION", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter", "Detail"], [
        ("Aircraft Type", ac["type"]),
        ("MSN", ac["msn"]),
        ("Registration", ac["reg"]),
        ("Year", str(ac["year"])),
        ("Engine Type", ac["engines"]),
        ("Engine Serial Numbers", f"{ac['eng_esn'][0]} (Pos.1) / {ac['eng_esn'][1]} (Pos.2)"),
        ("Current Total Time (approx.)", "30,000 FH / 21,500 FC (as of 1 Feb 2026)"),
        ("Interior Config.", "168 pax — to be confirmed per inspection"),
        ("Proposed Delivery Location", "Dallas/Fort Worth International (DFW)"),
        ("Proposed Delivery Date", "1 October 2026"),
    ], [2.2, 4.3])

    heading(doc, "2. COMMERCIAL TERMS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Term", "Proposed Value", "Notes"], [
        ("Lease Term", "84 months (7 years)", "Subject to ALA negotiation"),
        ("Commencement", "1 October 2026", "Subject to delivery conditions"),
        ("Expiry", "30 September 2033", "—"),
        ("Monthly Base Rent", f"USD {int(ac['monthly_rent']*1.08):,}", "Based on current appraisal; to be confirmed"),
        ("Rent Escalation", "CPI-U annually; floor 2%, ceiling 5%", "—"),
        ("Security Deposit", f"USD {int(ac['monthly_rent']*1.08*3):,} (3 months)", "Cash or acceptable L/C"),
        ("Maintenance Reserves", "Per Schedule 3 to ALA — rates TBD", "Subject to technical inspection"),
        ("Delivery Condition", "Half-life or better on all major components", "To be confirmed via PDI"),
        ("Governing Law", "New York / ICC Arbitration", "Standard ALA terms"),
        ("Exclusivity Period", "45 days from LOI execution", "Lessor commits not to market to third parties"),
        ("Non-Refundable Deposit", "USD 250,000 payable on LOI execution", "Credited to Security Deposit at lease start"),
    ], [1.6,1.8,3.1])

    page_break(doc)
    heading(doc, "3. MR TERM SHEET", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component","Proposed Rate","Escalation"], [
        ("Airframe A/B Check", f"USD {int(ac['mr_rate_airframe']*1.08)}/FH", "CPI-U"),
        ("Airframe C Check", f"USD {int(ac['mr_rate_airframe']*3*1.08)}/FH", "CPI-U"),
        ("Engine LLP – No.1", f"USD {int(ac['mr_rate_llp']*1.08)}/FC", "CPI-U + 1%"),
        ("Engine LLP – No.2", f"USD {int(ac['mr_rate_llp']*1.08)}/FC", "CPI-U + 1%"),
        ("Engine Performance – No.1", f"USD {int(ac['mr_rate_engine']*1.08)}/FH", "CPI-U + 1%"),
        ("Engine Performance – No.2", f"USD {int(ac['mr_rate_engine']*1.08)}/FH", "CPI-U + 1%"),
        ("Landing Gear", f"USD {int(ac['mr_rate_lg']*1.08)}/FH", "CPI-U"),
        ("APU", f"USD {int(ac['mr_rate_apu']*1.08)}/APU-H", "CPI-U"),
    ], [2.2, 1.8, 2.5])

    heading(doc, "4. CONDITIONS PRECEDENT TO ALA EXECUTION", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref","Condition"], [
        ("CP-01","Satisfactory pre-delivery inspection (physical, technical, borescope)"),
        ("CP-02","Satisfactory financial due diligence (audited accounts last 2 fiscal years)"),
        ("CP-03","FAA IASA Category 1 status maintained throughout negotiations"),
        ("CP-04","Insurance brokers confirm coverages placeable on Lessor's standard terms"),
        ("CP-05","Current IOSA certification (within 24 months)"),
        ("CP-06","No Material Adverse Change in Lessee's condition"),
        ("CP-07","Satisfactory legal opinion from Lessee's external counsel"),
        ("CP-08","No existing liens or encumbrances on Aircraft"),
        ("CP-09","Agreement on final form of ALA and all schedules"),
        ("CP-10","Payment of Non-Refundable Deposit"),
    ], [0.6, 5.9])

    heading(doc, "5. EXCLUSIVITY, CONFIDENTIALITY, AND EXPIRY", 2); add_horiz_rule(doc)
    body(doc, "5.1  Exclusivity. Upon execution and Deposit payment, Lessor commits to a 45-day exclusivity period during which it will not enter into any agreement with any third party regarding the Aircraft.")
    body(doc, "5.2  Confidentiality. Both Parties keep terms confidential; disclosure only to advisors on a need-to-know basis.")
    body(doc, "5.3  Non-Binding. Except for Deposit, exclusivity, and confidentiality provisions, this LOI is not legally binding until ALA execution.")
    body(doc, "5.4  Expiry. This LOI expires automatically if the ALA is not executed within 45 days of execution of this LOI.")

    sig_block(doc, [
        ("PROSPECTIVE LESSOR", "Niamh Gallagher", "VP Asset Sales & Leasing", "2026-02-10"),
        ("PROSPECTIVE LESSEE", "Dr. Elena V. Kowalski", "CEO", "2026-02-10"),
    ])

    path = os.path.join(OUT, "06-letters-of-intent", f"LOI-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: LOI-{ac['reg']}-{ac['msn']}.docx")


def build_insurance_cert(ac, idx):
    doc = Document()
    ref = f"IC-{ac['reg']}-2025-01"
    add_header_footer(doc, ref, "Certificate of Insurance", ac["lessee"])
    add_cover_page(doc, "Certificate of Insurance", ref, ac,
                   "Aviation Insurance — Certificate of Cover",
                   [("Named Insured", ac["lessee"], "USA"),
                    ("Additional Insured / Loss Payee", ac["lessor"], "Ireland"),
                    ("Broker", "Gallagher Heath Aviation", "Lloyd's, London")],
                   status="VALID — NOT A POLICY DOCUMENT")

    heading(doc, "1. POLICYHOLDER AND PERIOD", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Field","Detail"], [
        ("Policyholder", ac["lessee"]),
        ("Additional Insured / Loss Payee", f"{ac['lessor']} and affiliates, successors, assigns"),
        ("Broker", "Gallagher Heath Aviation, 1 Lime Street, London EC3M 7HA"),
        ("Policy No.", f"AV-2025-{ac['msn'][-4:]}-BDX"),
        ("Lloyd's Policy Ref.", "B0711/AV/2025/001234"),
        ("Policy Period", "01 Jan 2025 — 31 Dec 2025 (12:01 AM local time each date)"),
        ("Aircraft on Cover", f"{ac['type']} / MSN {ac['msn']} / Reg {ac['reg']}"),
        ("Agreed Hull Value", "USD 52,000,000"),
    ], [2.2, 4.3])

    heading(doc, "2. COVERAGES", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Coverage","Limit","Form","Deductible","Key Terms"], [
        ("Hull All Risks (Ground & Flight)","USD 52,000,000 AV","AV1A + AVN52H","USD 250,000","Lessor loss payee; no co-insurance; subrogation waiver"),
        ("Hull War & Allied Perils","USD 52,000,000 AV","AVN52H","USD 500,000","Confiscation/requisition; 7 days' notice"),
        ("Third-Party Liability","USD 750M per occurrence","AVN1C","Nil","Lessor additional insured; cross-liability; severability"),
        ("Passenger Liability","USD 300K × 168 seats = USD 50.4M","Montreal Convention","Nil","Per ICAO Annex 9"),
        ("Baggage & Cargo","USD 15M per occurrence","Montreal Convention","Nil","—"),
        ("Workers' Comp / EL","Statutory + USD 10M EL","Standard endorsement","Nil","Subrogation waiver for Lessor"),
        ("Terrorism Insurance","USD 52M AV","AVN52H","USD 250,000","7 days' cancellation notice"),
    ], [1.5,1.2,0.9,0.9,2.0])

    page_break(doc)
    heading(doc, "3. SUBSCRIBING MARKET", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Syndicate / Company","Market","Participation","AM Best Rating"], [
        ("Syndicate 2623 (Beazley)","Lloyd's","15.0%","A / XV"),
        ("Syndicate 1886 (Markel)","Lloyd's","12.5%","A / XV"),
        ("Syndicate 4444 (Atrium)","Lloyd's","10.0%","A- / XIV"),
        ("Syndicate 382 (AMS)","Lloyd's","8.0%","A / XIV"),
        ("Syndicate 2791 (MAP)","Lloyd's","7.5%","A- / XIII"),
        ("AXA XL Insurance","Company","18.0%","A / XV"),
        ("AIG Aviation Insurance","Company","14.0%","A+ / XV"),
        ("Chubb Aviation","Company","10.0%","AA- / XV"),
        ("Allianz Aviation Underwriters","Company","5.0%","AA / XV"),
        ("TOTAL","","100%","—"),
    ], [2.2,1.5,1.2,1.6])

    heading(doc, "4. KEY ENDORSEMENTS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref","Endorsement","Effect"], [
        ("AV52H","War and Allied Perils extension","Extends Hull All Risks to war, sabotage, terrorism, confiscation"),
        ("AVN1C","Airline Legal Liability — combined single limit","Bodily injury, property damage, passenger combined USD 750M"),
        ("AVN51","Noise, Pollution, Other Perils exclusion","Standard exclusions; aircraft-generated noise/pollution excluded unless crash"),
        ("AVN67B","Smooth closure","Continuity of cover on renewal"),
        ("LSW555D","Lessor's Interest Endorsement","Lessor as additional insured and loss payee; waiver of subrogation"),
        ("AVN38C","Cancellation notice","30 days to Lessor; 7 days for war risk"),
        ("SANCXXX","Sanctions","Policy void for claims from sanctioned territories/persons"),
    ], [1.0,2.5,3.0])

    sig_block(doc, [
        ("ISSUING BROKER", "Amanda J. Forsyth", "Aviation Director, Gallagher Heath", "2025-01-02"),
        ("ACKNOWLEDGED BY LESSEE", "", "VP Risk Management", ""),
    ])

    path = os.path.join(OUT, "07-insurance-certificates", f"IC-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: IC-{ac['reg']}-{ac['msn']}.docx")


def build_tar(ac, idx):
    doc = Document()
    ref = f"TAR-{ac['reg']}-{ac['lease_start']}"
    add_header_footer(doc, ref, "Technical Acceptance Report", ac["lessor"])
    add_cover_page(doc, "Technical Acceptance Report", ref, ac,
                   f"Pre-Delivery Technical Inspection — {ac['reg']}",
                   [("Lessor Inspector", "Walsh Aviation Consulting Ltd", "Ireland"),
                    ("Lessee Inspector", ac["lessee"], "USA"),
                    ("Independent", "Aviation Analysts Group LLC", "USA")],
                   status="ACCEPTED — DELIVERY APPROVED")

    heading(doc, "1. INSPECTION TEAM", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Role","Name","Company","Qualifications","Scope"], [
        ("Lessor Technical Rep","Patrick Walsh","Walsh Aviation Consulting","AME B1+B2, 25 yrs Boeing/CFM","Overall acceptance, records sign-off"),
        ("Lessee Technical Rep","Dr. Marina Torres","VP Technical","PhD Aerospace, ex-Boeing","Lessee acceptance, squawk sign-off"),
        ("Independent Inspector","George Nakamura","Aviation Analysts Group","A&P/IA, ex-FAA DER","Neutral verification"),
        ("Engine Inspector","Sean Callaghan","CFM Authorized Service Center","CFM56 factory-trained, Part 145","Engine borescope, LLP verification"),
        ("Structural Inspector","Hui Zhang","AAG Structures","NDT Level III (UT/FPI/Eddy)","Structural, corrosion assessment"),
        ("Records Specialist","Aoife Brennan","Walsh Aviation Consulting","10 yrs records auditing","Technical records completeness"),
    ], [1.5,1.3,1.5,1.5,1.7])

    page_break(doc)
    heading(doc, "2. GROUND INSPECTION RESULTS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["#","Area","Finding","Status","Action"], [
        ("GI-01","Nose / Radome","Radome crack-free; paint chip 1.2 cm lower skirt — no structural significance","ACCEPTABLE","Touch-up"),
        ("GI-02","Flight Deck","All panels complete; co-pilot EFIS brightness slightly low","ACCEPTABLE","Adjust EFIS brightness"),
        ("GI-03","Forward Lavatory","All systems functional","PASS","None"),
        ("GI-04","Business Class","168 seats inspected; 3 defects noted","MINOR DEFECTS","See squawk list"),
        ("GI-05","Aft Galley","All ovens/coffeemakers functional","PASS","None"),
        ("GI-06","Emergency Equipment","All in-date; counts correct","PASS","None"),
        ("GI-07","Cargo Compartments","Lining intact; no corrosion; door seals good","PASS","None"),
        ("GI-08","Nose Landing Gear","No corrosion; microswitches OK; tyre 70% worn — replace pre-delivery","ACTION","Replace NLG tyres before delivery"),
        ("GI-09","Left MLG","No corrosion; brake wear 45%; tyre 55%","ACCEPTABLE","Replace per Lessee programme"),
        ("GI-10","Right MLG","No corrosion; brake wear 38%; tyre 60%","ACCEPTABLE","None"),
        ("GI-11","Left Wing Lower","Sealant aging J10-L; slight seepage cold soak test","ACTION","Re-seal J10-L before delivery"),
        ("GI-12","Left Wing Upper","Spoilers 6×: functional; ailerons: functional; no disbonds","PASS","None"),
        ("GI-13","Right Wing","As Left Wing — all satisfactory","PASS","None"),
        ("GI-14","Engine No.1 Nacelle","Fan cowl OK; inlet no FOD; reverser seals good","PASS","None"),
        ("GI-15","Engine No.2 Nacelle","As Engine No.1","PASS","None"),
        ("GI-16","APU Bay","No oil leaks; exhaust duct clear","PASS","None"),
        ("GI-17","Fuselage Upper","No cracks; 3 minor repairs within SRM limits","PASS","None"),
        ("GI-18","Fuselage Lower","2 SRM repairs complete; paperwork complete","PASS","None"),
        ("GI-19","Empennage","No delaminations; trim tab functional; all static wicks present","PASS","None"),
        ("GI-20","Exterior Paint","Condition ~85%; no primer visible; marks correct","PASS","None"),
    ], [0.4,1.4,2.5,0.9,1.3])

    page_break(doc)
    heading(doc, "3. BORESCOPE — ENGINE NO. 1 (ESN " + ac["eng_esn"][0] + ")", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Module","Finding","Limit","Status"], [
        ("Fan (24 blades)","All within AMM; 2 blades: leading edge micro-nicks within 0.005 in limit","0.010 in","PASS"),
        ("LPC Booster","All serviceable","Per AMM 72-21-00","PASS"),
        ("HPC Stg 3","CSN at 14,990/15,000 — mandatory replacement at next SV per LLP sheet","15,000 FC","NOTED — approaching limit"),
        ("HPC Stg 4–9","Serviceable; coating wear within AMM; Stg 5 at 62% life","Per AMM","PASS"),
        ("Combustion Liner","Slight hot spot burn 2.5 cm — within SRM limits","SRM 72-41-00","PASS"),
        ("HPT Stg 1 blades","3 blades tip erosion at 80% of AMM limit — monitor","AMM 72-31-00","MONITOR"),
        ("HPT Stg 2 blades","All within AMM limits","—","PASS"),
        ("LPT Stg 4 Nozzle","All within limits; oxidation present but allowable","SRM 72-53-00","PASS"),
        ("LPT Stg 4–7 blades","All serviceable","Per AMM","PASS"),
        ("No. 1 Bearing","Carbon seals serviceable; no carbon dust migration","Per AMM 72-51-00","PASS"),
    ], [1.5,3.0,1.2,0.8])

    heading(doc, "4. BORESCOPE — ENGINE NO. 2 (ESN " + ac["eng_esn"][1] + ")", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Module","Finding","Status"], [
        ("Fan","All blades within limits; no FOD","PASS"),
        ("LPC","All blades serviceable","PASS"),
        ("HPC Stg 3","CSN 14,786 — within limit, approaching","PASS / MONITOR"),
        ("Combustion","Clean burn pattern; within limits","PASS"),
        ("HPT","Stg 1 blade erosion <50% AMM limit","PASS"),
        ("LPT","All stages serviceable","PASS"),
        ("Bearings","No abnormalities","PASS"),
    ], [1.5,3.5,1.5])

    page_break(doc)
    heading(doc, "5. ACCEPTANCE FLIGHT", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter","Detail"], [
        ("Flight Date", ac["lease_start"]),
        ("Route","DFW–DAL–DFW (acceptance circuit, 55 min airborne)"),
        ("Commander","Capt. Tom O'Brien (ATP, B737) — Lessor; Capt. Alicia Fernandez (ATP) — Lessee"),
        ("Weather","VFR; wind 180/12 kts; vis 10+; temp 22°C"),
        ("Engine No.1 EGT margin (T/O derate 5%)","36°C — PASS (min 30°C)"),
        ("Engine No.2 EGT margin","40°C — PASS"),
        ("APU in flight","Shut down post engine start; restarted in cruise for bleed test — PASS"),
        ("Pressurisation","Normal; cabin alt 6,200 ft at FL360"),
        ("Fuel burn","Within 2% of WBM prediction"),
        ("Autoland","ILS RWY 18L DFW; Cat I autoland; rollout normal"),
        ("All systems","Hydraulics, electrical, avionics, ACARS, FMC, autopilot — all PASS"),
        ("Acceptance Decision","ACCEPTED — Delivery Acceptance Certificate signed at gate"),
    ], [2.5,4.0])

    heading(doc, "6. PRE-DELIVERY ACTION ITEMS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Item","Description","Responsibility","Due"], [
        ("PD-01","Replace NLG tyres (GI-08)","Lessor","By Delivery Date"),
        ("PD-02","Re-seal wing panel J10-L (GI-11)","Lessor","By Delivery Date"),
        ("PD-03","Cosmetic paint touch-up radome skirt (GI-01)","Lessor — optional","Delivery + 30 days"),
    ], [0.8,3.0,1.3,1.4])
    body(doc, "All cabin squawks (SQK-01 through SQK-05) were rectified prior to the acceptance flight and are confirmed closed.")

    sig_block(doc, [
        ("LESSOR INSPECTOR", "Patrick Walsh", "Senior Technical Inspector", ac["lease_start"]),
        ("LESSEE INSPECTOR", "Dr. Marina Torres", "VP Technical", ac["lease_start"]),
        ("INDEPENDENT", "George Nakamura", "Aviation Analysts Group", ac["lease_start"]),
    ])

    path = os.path.join(OUT, "08-technical-acceptance-reports", f"TAR-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: TAR-{ac['reg']}-{ac['msn']}.docx")


def build_default_notice(ac, idx):
    doc = Document()
    ref = f"NOD-{ac['reg']}-2025-{idx:02d}"
    add_header_footer(doc, ref, "Notice of Default", ac["lessor"])
    add_cover_page(doc, "Notice of Default and Demand for Remedy", ref, ac,
                   "NOTICE OF EVENT OF DEFAULT",
                   [("Sending Party (Lessor)", ac["lessor"], "Ireland"),
                    ("Receiving Party (Lessee)", ac["lessee"], "USA")],
                   status="WITHOUT PREJUDICE — STRICTLY CONFIDENTIAL")

    # Red warning banner
    warn = doc.add_paragraph()
    r_w = warn.add_run("  ⚠  THIS IS A FORMAL NOTICE OF DEFAULT REQUIRING IMMEDIATE ACTION  ⚠  ")
    r_w.font.name = "Calibri"; r_w.font.size = Pt(10); r_w.bold = True; r_w.font.color.rgb = WHITE
    pPr = warn._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"),"CC0000"); pPr.append(shd)
    warn.paragraph_format.space_after = Pt(12)

    body(doc, "Date: 2025-09-15")
    body(doc, f"To: {ac['lessee']}")
    body(doc, f"From: {ac['lessor']}")
    body(doc, f"Re: Lease Ref. ALA-{2024+idx:04d}-{idx:03d} / {ac['type']} MSN {ac['msn']} Reg {ac['reg']} — NOTICE OF EVENT OF DEFAULT")

    page_break(doc)
    heading(doc, "1. ARREARS OF BASE RENT AND MAINTENANCE RESERVES", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Month","Component","Due Date","Due (USD)","Received (USD)","Outstanding (USD)","Days Overdue"], [
        ("Jun 2025","Base Rent","2025-06-01",f"{ac['monthly_rent']:,}","0",f"{ac['monthly_rent']:,}","106"),
        ("Jun 2025","Engine Perf MR – Eng 1","2025-06-10",f"{ac['mr_rate_engine']*320:,}","0",f"{ac['mr_rate_engine']*320:,}","96"),
        ("Jun 2025","Engine Perf MR – Eng 2","2025-06-10",f"{ac['mr_rate_engine']*318:,}","0",f"{ac['mr_rate_engine']*318:,}","96"),
        ("Jun 2025","Engine LLP MR – Eng 1","2025-06-10",f"{ac['mr_rate_llp']*485:,}","0",f"{ac['mr_rate_llp']*485:,}","96"),
        ("Jun 2025","Engine LLP MR – Eng 2","2025-06-10",f"{ac['mr_rate_llp']*482:,}","0",f"{ac['mr_rate_llp']*482:,}","96"),
        ("Jul 2025","Base Rent","2025-07-01",f"{ac['monthly_rent']:,}","0",f"{ac['monthly_rent']:,}","76"),
        ("Jul 2025","All MR Components","2025-07-10","298,400","0","298,400","66"),
        ("Aug 2025","Base Rent","2025-08-01",f"{ac['monthly_rent']:,}","0",f"{ac['monthly_rent']:,}","45"),
        ("Aug 2025","All MR Components","2025-08-10","298,400","0","298,400","36"),
        ("Sep 2025","Base Rent","2025-09-01",f"{ac['monthly_rent']:,}","0",f"{ac['monthly_rent']:,}","14"),
    ], [0.8,1.5,0.9,0.9,0.9,0.9,0.8])
    body(doc, f"TOTAL ARREARS OF BASE RENT: USD {ac['monthly_rent']*4:,}")
    body(doc, "TOTAL ARREARS OF MR: USD 1,234,560 (to be confirmed)")
    body(doc, "LATE PAYMENT INTEREST (SOFR + 4.00%): USD 38,412")
    body(doc, f"TOTAL OUTSTANDING: APPROXIMATELY USD {ac['monthly_rent']*4 + 1272972:,}")

    heading(doc, "2. INSURANCE DEFAULT", 2); add_horiz_rule(doc)
    body(doc, "Lessor has been advised that the Hull All Risks policy is subject to cancellation on 30 September 2025 for non-payment of premium (outstanding: USD 842,000). This constitutes an Event of Default per Article 20.1(c). DEMAND: Lessee must provide evidence of reinstatement by 17:00 NY time on 20 September 2025.")

    heading(doc, "3. MAINTENANCE DEFAULT", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Ref","ATA","Default","Required Action","Deadline"], [
        ("MDEF-01","72","Engine No.1 EGT margin at 17°C against lease minimum of 20°C — immediate shop visit planning required","Submit engine shop visit plan to Lessor within 10 Business Days","2025-09-29"),
        ("MDEF-02","05","A-Check overdue by 12 days (last performed 601 FH ago against 600 FH interval)","Ground Aircraft immediately; perform A-Check; advise completion","Immediately"),
        ("MDEF-03","31","AD 2025-18-004 TCAS II 7.1 upgrade (compliance required 31 Aug 2025) NOT INSTALLED","Comply with AD immediately; provide proof of compliance","Immediately"),
        ("MDEF-04","32","NLG shock strut servicing overdue (interval 500 FH; last at 684 FH ago)","Perform NLG shock strut service","2025-09-18"),
    ], [0.7,0.6,2.5,1.6,0.9])

    page_break(doc)
    heading(doc, "4. LESSEE RESPONSE REQUIRED BY 22 SEPTEMBER 2025", 2); add_horiz_rule(doc)
    body(doc, "Within five (5) Business Days of this Notice, Lessee must:")
    body(doc, "(a)  Pay all outstanding arrears in full together with all accrued interest;", indent=True)
    body(doc, "(b)  Provide evidence of reinstatement of required insurance coverages;", indent=True)
    body(doc, "(c)  Provide a written remediation plan for each maintenance default in Section 3; and", indent=True)
    body(doc, "(d)  Provide a written explanation of the circumstances leading to these Events of Default.", indent=True)

    heading(doc, "5. RESERVATION OF RIGHTS", 2); add_horiz_rule(doc)
    body(doc, "Nothing in this Notice constitutes a waiver of any right or remedy. Lessor expressly reserves all rights including: (a) terminate the Lease with immediate effect; (b) repossess the Aircraft; (c) draw upon the Security Deposit; (d) claim all amounts due plus damages for the unexpired Lease Period; and (e) seek any other remedy available at law.")

    body(doc, "")
    body(doc, f"For and on behalf of {ac['lessor']}")
    body(doc, "")
    body(doc, "________________________")
    body(doc, "Head of Portfolio Management")
    body(doc, ac["lessor"])

    path = os.path.join(OUT, "09-default-notices", f"NOD-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: NOD-{ac['reg']}-{ac['msn']}.docx")


def build_supp_rent(ac, idx):
    doc = Document()
    ref = f"SRS-{ac['reg']}-Q3-2025"
    add_header_footer(doc, ref, "Supplemental Rent Statement", ac["lessor"])
    add_cover_page(doc, "Supplemental Rent Statement", ref, ac,
                   "Q3 2025 — Base Rent & Maintenance Reserves",
                   [("Lessor", ac["lessor"], "Ireland"),
                    ("Lessee", ac["lessee"], "USA")],
                   status="INVOICE — PAYMENT DUE 20 OCT 2025")

    heading(doc, "1. STATEMENT DETAILS", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Parameter","Detail"], [
        ("Statement Reference", ref),
        ("Aircraft", f"{ac['type']} / {ac['reg']} / MSN {ac['msn']}"),
        ("Statement Period", "Q3 2025 (01 July – 30 September 2025)"),
        ("Statement Date", "2025-10-10"),
        ("Payment Due", "2025-10-20"),
        ("Lessor", ac["lessor"]),
        ("Lessee", ac["lessee"]),
        ("Lease Reference", f"ALA-{2024+idx:04d}-{idx:03d}"),
    ], [2.2,4.3])

    heading(doc, "2. UTILISATION — Q3 2025", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Month","Flight Hours","Flight Cycles","APU Hours","Sectors"], [
        ("July 2025","482","341","68","341"),
        ("August 2025","511","362","74","362"),
        ("September 2025","476","337","71","337"),
        ("Q3 2025 TOTAL","1,469","1,040","213","1,040"),
        ("Cumulative (since delivery)","14,822","10,504","2,184","10,504"),
    ], [1.8,1.4,1.4,1.4,1.5])

    page_break(doc)
    heading(doc, "3. MAINTENANCE RESERVE CALCULATION — Q3 2025", 2); add_horiz_rule(doc)
    q3_mr = (ac['mr_rate_airframe']*1469 + ac['mr_rate_airframe']*3*1469 +
             ac['mr_rate_engine']*1469*2 + ac['mr_rate_llp']*1040*2 +
             ac['mr_rate_lg']*1469 + ac['mr_rate_apu']*213)
    add_table_styled(doc, ["Reserve Component","Rate","Q3 Utilisation","Q3 Amount (USD)","YTD Cumulative (USD)"], [
        ("Airframe A/B Check", f"USD {ac['mr_rate_airframe']}/FH","1,469 FH", f"{ac['mr_rate_airframe']*1469:,}", f"{ac['mr_rate_airframe']*4432:,}"),
        ("Airframe C Check", f"USD {ac['mr_rate_airframe']*3}/FH","1,469 FH", f"{ac['mr_rate_airframe']*3*1469:,}", f"{ac['mr_rate_airframe']*3*4432:,}"),
        ("Engine Performance – No.1", f"USD {ac['mr_rate_engine']}/FH","1,469 FH", f"{ac['mr_rate_engine']*1469:,}", f"{ac['mr_rate_engine']*4432:,}"),
        ("Engine Performance – No.2", f"USD {ac['mr_rate_engine']}/FH","1,469 FH", f"{ac['mr_rate_engine']*1469:,}", f"{ac['mr_rate_engine']*4432:,}"),
        ("Engine LLP – No.1", f"USD {ac['mr_rate_llp']}/FC","1,040 FC", f"{ac['mr_rate_llp']*1040:,}", f"{ac['mr_rate_llp']*3133:,}"),
        ("Engine LLP – No.2", f"USD {ac['mr_rate_llp']}/FC","1,040 FC", f"{ac['mr_rate_llp']*1040:,}", f"{ac['mr_rate_llp']*3133:,}"),
        ("Landing Gear (all legs)", f"USD {ac['mr_rate_lg']}/FH","1,469 FH", f"{ac['mr_rate_lg']*1469:,}", f"{ac['mr_rate_lg']*4432:,}"),
        ("APU", f"USD {ac['mr_rate_apu']}/APU-H","213 APU-H", f"{ac['mr_rate_apu']*213:,}", f"{ac['mr_rate_apu']*633:,}"),
        ("Q3 2025 TOTAL MR DUE","—","—", f"{q3_mr:,}", "—"),
    ], [2.0,1.1,1.2,1.2,1.8])

    heading(doc, "4. COMPONENT LIFE STATUS — 30 SEPT 2025", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Component","Life Limit","TSN at 30 Sep","Remaining","% Used","Alert Level"], [
        (f"HPC Disc Stg 3 – Eng 1 ({ac['eng_esn'][0]})","15,000 FC","14,822 FC","178 FC","98.8%","URGENT — within 2 months"),
        (f"HPC Disc Stg 3 – Eng 2 ({ac['eng_esn'][1]})","15,000 FC","14,618 FC","382 FC","97.5%","Within 4 months"),
        ("Fan Blades – Eng 1 (set)","20,000 FC","14,822 FC","5,178 FC","74.1%","Normal"),
        ("HPT Stg 1 Disc – Eng 1","15,000 FC","14,822 FC","178 FC","98.8%","URGENT — same as HPC Stg 3"),
        ("HPT Stg 1 Blade – Eng 1","10,000 FC (post-PRSV)","2,500 FC","7,500 FC","25.0%","Normal"),
        ("Airframe C Check","6,000 FH / 36 months","4,823 FH since C-Chk","1,177 FH","80.4%","Within 13 months"),
        ("NLG Overhaul","10 years","6 yrs since OH","4 yrs remaining","60%","Normal"),
        ("MLG Overhaul","10 years","5.5 yrs since OH","4.5 yrs remaining","55%","Normal"),
        ("APU HSI","4,000 APU-H","2,184 APU-H since HSI","1,816 APU-H","54.6%","Normal"),
    ], [2.0,1.0,1.3,1.1,0.8,1.3])

    page_break(doc)
    heading(doc, "5. PAYMENT SUMMARY", 2); add_horiz_rule(doc)
    add_table_styled(doc, ["Item","Amount (USD)"], [
        ("Base Rent — July 2025", f"USD {ac['monthly_rent']:,}"),
        ("Base Rent — August 2025", f"USD {ac['monthly_rent']:,}"),
        ("Base Rent — September 2025", f"USD {ac['monthly_rent']:,}"),
        ("Q3 2025 Maintenance Reserves (all components)", f"USD {q3_mr:,}"),
        ("Late payment interest (Art. 6.4)", "USD 0 (assumed timely)"),
        ("TOTAL PAYABLE BY 20 OCTOBER 2025", f"USD {ac['monthly_rent']*3 + q3_mr:,}"),
    ], [3.5,3.0])
    body(doc, f"Wire transfer instructions per Article 6.3 of Lease {ref.split('-')[0]}. Payment reference: {ref}")

    path = os.path.join(OUT, "10-supplemental-rent-statements", f"SRS-{ac['reg']}-{ac['msn']}.docx")
    doc.save(path)
    print(f"    Saved: SRS-{ac['reg']}-{ac['msn']}.docx")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print(f"Generating 45 fully-formatted aircraft leasing documents into: {OUT}")
    print("Features: cover pages, company logos, headers/footers, color tables, page numbers")
    print()

    for i, ac in enumerate(AIRCRAFT):
        print(f"Aircraft {i+1}/{len(AIRCRAFT)}: {ac['reg']} ({ac['type']})")
        print(f"  [1/5] Lease Agreement...")
        build_lease_agreement(ac, i)
        print(f"  [2/5] Delivery Condition Report...")
        build_dcr(ac, i)
        print(f"  [3/5] MR Claim...")
        build_mr_claim(ac, i)
        print(f"  [4/5] Return Condition Report...")
        build_rcr(ac, i)
        print(f"  [5/5] Lease Amendment...")
        build_amendment(ac, i)
        print()

    print("Additional document types...")
    for i in range(3):
        print(f"  LOI {i+1}: {AIRCRAFT[i]['reg']}...")
        build_loi(AIRCRAFT[i], i)
    for i in range(3):
        print(f"  Insurance Certificate {i+1}: {AIRCRAFT[i+2]['reg']}...")
        build_insurance_cert(AIRCRAFT[i+2], i)
    for i in range(3):
        print(f"  TAR {i+1}: {AIRCRAFT[i+1]['reg']}...")
        build_tar(AIRCRAFT[i+1], i)
    for i in range(3):
        print(f"  Default Notice {i+1}: {AIRCRAFT[i+3]['reg']}...")
        build_default_notice(AIRCRAFT[i+3], i)
    for i in range(3):
        print(f"  Supplemental Rent Statement {i+1}: {AIRCRAFT[i+1]['reg']}...")
        build_supp_rent(AIRCRAFT[i+1], i)

    print()
    print("=== RESULTS ===")
    total_size = 0
    for folder in sorted(os.listdir(OUT)):
        folder_path = os.path.join(OUT, folder)
        if os.path.isdir(folder_path):
            files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
            folder_size = sum(os.path.getsize(os.path.join(folder_path, f)) for f in files) // 1024
            total_size += folder_size
            print(f"  {folder}/  — {len(files)} files, {folder_size} KB")
    print(f"\nTotal: {total_size} KB across 45 documents")


if __name__ == "__main__":
    main()
